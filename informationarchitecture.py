#!/usr/bin/env python3
"""
==============================================================
 DBIM Information Architecture Scanner
 Focus: DBIM Figure 45 & Table 11 + UX Evidence Gathering
==============================================================
"""

import sys
import os
import time
import logging
import argparse
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import TimeoutException, NoSuchElementException, WebDriverException, StaleElementReferenceException
from webdriver_manager.chrome import ChromeDriverManager

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── LOGGING ───────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s", datefmt="%H:%M:%S")
log = logging.getLogger("dbim-ia-scanner")

# ── DBIM MANDATED ARCHITECTURE DICTIONARIES ───────────────────────────────────
DBIM_L1_MANDATORY = [
    ["home"],
    ["ministry", "department"],
    ["offerings"],
    ["documents"],
    ["media", "resources"],
    ["connect"]
]

DBIM_L2_EXPECTED = {
    "ministry_department": [
        "about us", "our team", "our organizations", "our groups", 
        "our performance", "directory"
    ],
    "offerings": [
        "schemes", "services", "competitions", "awards", 
        "international collaborations", "careers", "tenders"
    ],
    "documents": [
        "reports", "acts", "policies", "orders", "notices", 
        "publications", "press release", "gazette notifications"
    ],
    "resources_media": [
        "photos", "videos", "brochures", "presentations", "podcasts"
    ],
    "connect": [
        "contact us", "directory", "rti", "visitor's pass", "visitor pass",
        "grievance", "citizen engagement", "parliament questions"
    ]
}

SCREENSHOT_DIR = "screenshots"
REPORT_DOCX = "DBIM_IA_Compliance_Report.docx"

# ── WEBDRIVER SETUP ───────────────────────────────────────────────────────────
def make_driver(headless: bool = True) -> webdriver.Chrome:
    opts = Options()
    if headless:
        opts.add_argument("--headless=new")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--window-size=1920,1080")
    opts.add_argument("--log-level=3")
    opts.page_load_strategy = "eager" 
    opts.add_argument(
        "--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )
    opts.add_experimental_option("excludeSwitches", ["enable-logging"])

    try:
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=opts)
    except Exception as exc:
        log.warning("ChromeDriverManager failed (%s); falling back to PATH chromedriver.", exc)
        driver = webdriver.Chrome(options=opts)

    driver.set_page_load_timeout(15)
    return driver

# ── SCANNER CLASS ─────────────────────────────────────────────────────────────
class DBIMArchitectureScanner:
    def __init__(self, target_url: str, headless: bool = True):
        self.target_url = target_url
        self.driver = make_driver(headless)
        self.results = {}
        self.screenshot_path = ""
        log.info("Scanner initialised for DBIM IA validation → %s", self.target_url)

    def wait_for_page_load(self, timeout=15):
        try:
            WebDriverWait(self.driver, timeout).until(
                lambda d: d.execute_script("return document.readyState") == "complete"
            )
            time.sleep(2) 
        except TimeoutException:
            log.warning("Page load timeout reached. Proceeding with available DOM.")

    def take_screenshot(self, name: str) -> str:
        os.makedirs(SCREENSHOT_DIR, exist_ok=True)
        path = os.path.join(SCREENSHOT_DIR, f"{name}_{int(time.time())}.png")
        try:
            # Scroll to top before taking evidence screenshot
            self.driver.execute_script("window.scrollTo(0, 0);")
            time.sleep(0.5)
            self.driver.save_screenshot(path)
            return path
        except Exception as e:
            log.error("Failed to capture screenshot: %s", e)
            return ""

    def extract_menu_text(self) -> str:
        nav_selectors = [
            "nav", "#main-nav", "#menu", ".navbar", ".navigation", 
            "[role='navigation']", "header"
        ]
        for sel in nav_selectors:
            try:
                el = self.driver.find_element(By.CSS_SELECTOR, sel)
                if el.is_displayed():
                    return el.text.lower()
            except NoSuchElementException:
                continue
        return self.driver.find_element(By.TAG_NAME, "body").text.lower()

    def gather_ux_evidence(self):
        """Navigates site to gather screenshots for manual Consistency/Clarity checks."""
        log.info("Gathering Visual Evidence for Subjective UX Objectives...")
        evidence_records = []

        # 1. Homepage Baseline
        hp_ss = self.take_screenshot("UX_Evidence_Homepage_Baseline")
        evidence_records.append({
            "type": "Homepage Baseline",
            "url": self.driver.current_url,
            "screenshot": hp_ss,
            "instruction": "Review for: Simplification (cognitive load) and Engagement."
        })

        # 2. Interior Page Template
        try:
            # Find a valid internal link in the navigation
            nav_links = self.driver.find_elements(By.CSS_SELECTOR, "nav a, .navbar a, header a")
            for link in nav_links:
                try:
                    href = link.get_attribute("href")
                    if link.is_displayed() and href and href.startswith("http") and self.target_url in href and href != self.target_url and href != self.target_url + "/":
                        log.info(f"Navigating to interior page for Consistency check: {href}")
                        self.driver.get(href)
                        self.wait_for_page_load()
                        int_ss = self.take_screenshot("UX_Evidence_Interior_Template")
                        evidence_records.append({
                            "type": "Interior Template",
                            "url": self.driver.current_url,
                            "screenshot": int_ss,
                            "instruction": "Review for: Consistency. Check if layout, fonts, and headers match the Homepage baseline."
                        })
                        break # Only need one good interior example
                except StaleElementReferenceException:
                    continue
        except Exception as e:
            log.warning(f"Error gathering interior page evidence: {e}")
            
        self.results["ux_evidence"] = evidence_records

    def run_dbim_audit(self):
        log.info("Accessing target URL...")
        try:
            self.driver.get(self.target_url)
            self.wait_for_page_load()
        except TimeoutException:
            log.error("Failed to connect to %s within timeout boundaries.", self.target_url)
            self.driver.quit()
            return

        self.screenshot_path = self.take_screenshot("TC_Evidence_Navigation")
        nav_text = self.extract_menu_text()
        
        # Evaluate Level 1 Constraints
        log.info("Executing Test Case: Level 1 Global Menu Sections...")
        l1_results = []
        for index, required_group in enumerate(DBIM_L1_MANDATORY):
            found = any(kw in nav_text for kw in required_group)
            label = " / ".join(required_group).title()
            l1_results.append({
                "tc_id": f"TC-L1-0{index+1}",
                "section": label,
                "status": "PASS" if found else "FAIL",
                "finding": f"'{label}' found in global navigation." if found else f"Missing mandatory Level 1 section: '{label}'."
            })
        self.results["level_1"] = l1_results

        # Evaluate Level 2 Constraints
        log.info("Executing Test Case: Level 2 Illustrative Groupings...")
        l2_results = []
        for index, (category, keywords) in enumerate(DBIM_L2_EXPECTED.items()):
            category_label = category.replace("_", " ").title()
            found_items = [kw for kw in keywords if kw in nav_text]
            missing_items = [kw for kw in keywords if kw not in nav_text]
            
            if missing_items:
                status, finding = "WARN", f"Found {len(found_items)} items. Missing suggested L2 items: {', '.join(missing_items).title()}."
            else:
                status, finding = "PASS", f"All illustrative Level 2 items found for {category_label}."
                
            l2_results.append({
                "tc_id": f"TC-L2-0{index+1}",
                "section": category_label,
                "status": status,
                "finding": finding
            })
        self.results["level_2"] = l2_results

        # Gather UX Evidence
        self.gather_ux_evidence()

        self.driver.quit()
        self.generate_docx_report()
        self.print_console_summary()


    # ── REPORTING ─────────────────────────────────────────────────────────────
    def embed_image(self, doc, path, title=None):
        if path and os.path.exists(path):
            if title:
                doc.add_paragraph(title).runs[0].bold = True
            try:
                doc.add_picture(path, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f"[Failed to embed screenshot: {e}]")
            doc.add_paragraph("")

    def generate_docx_report(self):
        if not HAS_DOCX: return

        STATUS_COLORS = {"PASS": RGBColor(0x27, 0xAE, 0x60), "WARN": RGBColor(0xE6, 0x7E, 0x22), "FAIL": RGBColor(0xE7, 0x4C, 0x3C)}
        doc = DocxDocument()
        
        # Header
        doc.add_heading("DBIM Information Architecture Test Report", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Target Website: {self.target_url}\n").bold = True
        meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_page_break()

        # Level 1
        doc.add_heading("Test Case 1: Mandatory Global Navigation (Figure 45)", level=1)
        t1 = doc.add_table(rows=1, cols=4)
        t1.style = "Table Grid"
        h1 = t1.rows[0].cells
        h1[0].text, h1[1].text, h1[2].text, h1[3].text = "Test Case ID", "Target Section", "Status", "Reason / Finding"
        for c in h1: c.paragraphs[0].runs[0].bold = True
        
        for item in self.results.get("level_1", []):
            row = t1.add_row().cells
            row[0].text, row[1].text, row[3].text = item["tc_id"], item["section"], item["finding"]
            status_run = row[2].paragraphs[0].add_run(item["status"])
            status_run.bold, status_run.font.color.rgb = True, STATUS_COLORS[item["status"]]

        doc.add_paragraph("")
        self.embed_image(doc, self.screenshot_path, "Test Evidence (Navigation Visual Capture):")
        doc.add_page_break()

        # Level 2
        doc.add_heading("Test Case 2: Content Grouping Validation (Table 11)", level=1)
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = "Table Grid"
        h2 = t2.rows[0].cells
        h2[0].text, h2[1].text, h2[2].text, h2[3].text = "Test Case ID", "Parent Category", "Status", "Reason / Finding"
        for c in h2: c.paragraphs[0].runs[0].bold = True

        for item in self.results.get("level_2", []):
            row = t2.add_row().cells
            row[0].text, row[1].text, row[3].text = item["tc_id"], item["section"], item["finding"]
            status_run = row[2].paragraphs[0].add_run(item["status"])
            status_run.bold, status_run.font.color.rgb = True, STATUS_COLORS[item["status"]]

        doc.add_paragraph("")
        self.embed_image(doc, self.screenshot_path, "Test Evidence (Navigation Visual Capture):")
        doc.add_page_break()

        # UX Evidence Section
        doc.add_heading("Test Case 3: Subjective UX Objectives Evidence", level=1)
        doc.add_paragraph("Objective: Gather visual evidence to manually assess DBIM heuristics for Simplification, Consistency, and Engagement. Automation cannot assign pass/fail metrics to subjective design principles.")
        
        for record in self.results.get("ux_evidence", []):
            doc.add_heading(f"Evidence Type: {record['type']}", level=2)
            doc.add_paragraph(f"URL: {record['url']}")
            doc.add_paragraph(f"Review Instruction: {record['instruction']}").runs[0].bold = True
            self.embed_image(doc, record["screenshot"])

        doc.save(REPORT_DOCX)
        log.info("DOCX Report saved -> %s", REPORT_DOCX)

    def print_console_summary(self):
        ICONS = {"PASS": "✅", "WARN": "⚠️ ", "FAIL": "❌"}
        print("\n" + "=" * 100)
        print(f"  DBIM INFORMATION ARCHITECTURE TEST SUMMARY: {self.target_url}")
        print("=" * 100)
        
        print("\n[ TEST CASE 1: Level 1 Mandatory Sections ]")
        for item in self.results.get("level_1", []):
            print(f"  {ICONS[item['status']]} {item['tc_id']} | {item['section']:<25} → {item['finding']}")

        print("\n[ TEST CASE 2: Level 2 Illustrative Groupings ]")
        for item in self.results.get("level_2", []):
            print(f"  {ICONS[item['status']]} {item['tc_id']} | {item['section']:<25} → {item['finding']}")

        print("\n[ TEST CASE 3: UX Objectives Evidence Gathered ]")
        for record in self.results.get("ux_evidence", []):
            print(f"  📸 Captured: {record['type']:<20} | URL: {record['url']}")
                
        print("\n" + "=" * 100 + "\n")


# ── ENTRY POINT ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("url", nargs="?", default=None)
    parser.add_argument("--no-headless", action="store_true")
    args = parser.parse_args()

    target_url = args.url or input("\n🌐 Enter website URL to execute DBIM IA Test Cases: ").strip()
    if not target_url: sys.exit(1)
    if not target_url.startswith(("http://", "https://")): target_url = "https://" + target_url

    DBIMArchitectureScanner(target_url, headless=not args.no_headless).run_dbim_audit()