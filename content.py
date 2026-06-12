import os
import re
import time
import shutil
import argparse
import requests
import warnings

from datetime import datetime
from urllib.parse import urlparse, urljoin

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import StaleElementReferenceException

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

warnings.filterwarnings("ignore", message="Unverified HTTPS request")


# =============================================================================
# OPTIONAL IMPORTS
# =============================================================================

try:
    from spellchecker import SpellChecker
    _SPELL_OK = True
except ImportError:
    _SPELL_OK = False
    print(
        "[WARN] pyspellchecker not installed. "
        "Falling back to basic error list.\n"
        "       Fix: pip install pyspellchecker"
    )

try:
    import language_tool_python
    _GRAMMAR_OK = True
except ImportError:
    _GRAMMAR_OK = False
    print(
        "[WARN] language_tool_python not installed. "
        "Grammar check will be skipped.\n"
        "       Fix: pip install language-tool-python"
    )


# =============================================================================
# CONSTANTS — EXISTING
# =============================================================================

HINGLISH_WORDS = [
    "hai", "haan", "nahi", "kya", "kyun",
    "kaise", "kab", "kahan", "bhi",
    "toh", "sirf", "lekin", "aur",
    "ya", "par", "mein", "aap",
    "hum", "yeh", "woh"
]

AMERICAN_SPELLINGS = [
    "color", "flavor", "honor", "neighbor",
    "organize", "recognize", "analyze",
    "center", "fiber", "theater",
    "traveled", "canceled", "program"
]


# =============================================================================
# CONSTANTS — CL-05: Content Freshness (A.5.6)
# =============================================================================

STALE_THRESHOLD_MONTHS = 24

LAST_UPDATED_PATTERNS = [
    r"last\s+updated?\s*[:\-]?\s*(\d{1,2}[\s\-/]\w+[\s\-/]\d{2,4})",
    r"updated\s+on\s*[:\-]?\s*(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})",
    r"date\s*[:\-]\s*(\d{4}[/\-\.]\d{1,2}[/\-\.]\d{1,2})",
    r"(\d{4}-\d{2}-\d{2})",
]

PLACEHOLDER_PATTERNS = [
    r"\blorem\s+ipsum\b",
    r"\bcoming\s+soon\b",
    r"\bunder\s+construction\b",
    r"\bplaceholder\b",
    r"\bto\s+be\s+updated\b",
    r"\btbd\b",
    r"\bfixme\b",
    r"\binsert\s+content\s+here\b",
    r"\bsample\s+text\b",
    r"\bdummy\s+text\b",
    r"\bcontent\s+here\b",
    r"\benter\s+text\s+here\b",
]

CONTENT_SECTION_SELECTORS = [
    "main", "article", "section",
    ".content", ".main-content", "#content",
    ".page-content", "#main-content"
]

MAX_INTERNAL_LINKS_TO_CHECK = 20


# =============================================================================
# CONSTANTS — CL-06: Grammar & Spelling Validation (7.1.3.3)
# =============================================================================

SPELL_MIN_WORD_LEN    = 4
SPELL_MAX_SAMPLE      = 10
EXCESSIVE_CAP_MIN_LEN = 5

SPELL_WHITELIST = {
    "govt", "gov", "nic", "india", "hindi", "dbim", "ccps",
    "url", "http", "https", "pdf", "api", "ui", "ux", "id",
    "jan", "feb", "mar", "apr", "jun", "jul", "aug",
    "sep", "oct", "nov", "dec", "lakh", "crore",
    "rupee", "rupees", "digitalindia", "egovernance",
    "meity", "miety", "uidai", "aadhaar", "aadhar",
    "umang", "digilocker", "mygov", "bharatnet",
    "nicnet", "niccloud", "meghraj", "stqc",
    "goi", "mha", "mea", "ias", "ips", "rti",
    "upsc", "ssb", "ndma", "cic", "pmo",
    "niti", "aayog", "darpg", "doit", "deity",
    "nsdg", "negp", "csc", "sdc", "swan",
    "nkn", "aebas", "pfms", "cpgrams", "saral",
    "otp", "sms", "qr", "uuid", "html", "css",
    "json", "xml", "rest", "soap", "sdk",
    "oauth", "saml", "ldap", "cms", "crm",
    "erp", "sso", "cdn", "dns", "ssl", "tls",
    "smtp", "ftp", "ssh", "vpc", "aws", "gcp",
    "captcha", "ajax", "jquery", "nodejs",
    "pradesh", "bharat", "dilli", "mumbai",
    "bengaluru", "kolkata", "chennai",
    "hyderabad", "ahmedabad", "pune",
    "panchayat", "zila", "tehsil", "taluk",
    "viz", "etc", "approx", "dept", "min", "max",
    "no", "nos", "sr", "mr", "mrs", "dr", "prof",
    "avg", "std", "vol", "pg", "pp", "ed", "rev",
    "colour", "flavour", "honour", "neighbour",
    "organise", "recognise", "analyse",
    "centre", "fibre", "theatre",
    "travelled", "cancelled", "programme",
    "defence", "licence", "practise",
    "specialise", "realise", "utilise",
}

KNOWN_CAPS_TOKENS = {
    "HTTPS", "HTML", "URL", "NIC", "DBIM", "CCPS", "STQC",
    "FAQ", "PDF", "API", "INDIA", "PASS", "FAIL", "SKIP",
    "GOI", "MHA", "MEA", "IAS", "IPS", "RTI", "OTP", "SMS",
    "MEITY", "UIDAI", "UMANG", "MYGOV", "UPSC", "NDMA",
    "NITI", "DARPG", "PFMS", "CPGRAMS", "JSON", "XML",
    "REST", "SOAP", "UUID", "SAML", "LDAP", "AJAX",
    "CSS", "SDK", "CDN", "DNS", "SSL", "TLS", "SSO",
    "CMS", "CRM", "ERP", "QR", "UX", "UI", "IT",
    "ICT", "NKN", "SDC", "CSC", "SWAN", "NEGP",
    "NSDG", "AAYOG", "PMO", "CIC", "SSB",
}

FALLBACK_MISSPELLINGS = [
    ("recieve",        "receive"),
    ("occured",        "occurred"),
    ("seperate",       "separate"),
    ("definately",     "definitely"),
    ("goverment",      "government"),
    ("accomodation",   "accommodation"),
    ("beleive",        "believe"),
    ("relevent",       "relevant"),
    ("publically",     "publicly"),
    ("writting",       "writing"),
    ("committment",    "commitment"),
    ("adress",         "address"),
    ("basicaly",       "basically"),
    ("correspondance", "correspondence"),
]

_RE_EMAIL = re.compile(
    r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
    re.IGNORECASE
)

_RE_URL = re.compile(
    r"https?://[^\s<>\"']+"
    r"|ftp://[^\s<>\"']+",
    re.IGNORECASE
)

_RE_DOMAIN = re.compile(
    r"\b(?:www\.)?"
    r"[a-zA-Z0-9\-]+"
    r"(?:\.[a-zA-Z0-9\-]+){1,4}"
    r"\."
    r"(?:com|org|net|in|gov|edu|co\.in|gov\.in|nic\.in|"
    r"ac\.in|res\.in|mil|int)\b",
    re.IGNORECASE
)

_RE_SNAKE_CASE = re.compile(
    r"\b[a-z][a-z0-9]*(?:_[a-z0-9]+)+\b"
)

_RE_CAMEL_CASE = re.compile(
    r"\b[a-zA-Z]+(?:[A-Z][a-z0-9]+)+\b"
)

_RE_KEBAB_CASE = re.compile(
    r"\b[a-z][a-z0-9]*(?:-[a-z0-9]+)+\b"
)

_RE_UUID = re.compile(
    r"\b[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}"
    r"-[0-9a-f]{4}-[0-9a-f]{12}\b",
    re.IGNORECASE
)

_RE_HASH = re.compile(
    r"\b[0-9a-f]{16,}\b",
    re.IGNORECASE
)

_RE_FILENAME = re.compile(
    r"\b[\w\-]+\."
    r"(?:pdf|doc|docx|xls|xlsx|ppt|pptx|"
    r"csv|txt|json|xml|html|htm|css|js|"
    r"png|jpg|jpeg|gif|svg|ico|mp4|mp3|"
    r"zip|rar|tar|gz)\b",
    re.IGNORECASE
)

_RE_API_PATH = re.compile(
    r"(?<!\w)/[a-zA-Z0-9_\-]+(?:/[a-zA-Z0-9_\-]+)+"
)

_RE_NUMERIC = re.compile(
    r"\b\d[\d\s\-().+]{3,}\b"
)

_RE_VERSION = re.compile(
    r"\bv?\d+\.\d+(?:\.\d+)*\b",
    re.IGNORECASE
)

_RE_CSS_CLASS = re.compile(
    r"\.[a-z][a-z0-9]*(?:-[a-z0-9]+)+",
    re.IGNORECASE
)

_EXCLUSION_PATTERNS = [
    _RE_EMAIL,
    _RE_URL,
    _RE_UUID,
    _RE_HASH,
    _RE_FILENAME,
    _RE_API_PATH,
    _RE_NUMERIC,
    _RE_VERSION,
    _RE_DOMAIN,
    _RE_SNAKE_CASE,
    _RE_CAMEL_CASE,
    _RE_KEBAB_CASE,
    _RE_CSS_CLASS,
]


# =============================================================================
# CONSTANTS — CCPS-01: CCPS Integration (7.4) — CONFIDENCE-BASED
# =============================================================================

CCPS_HIGH_CONFIDENCE = [
    "ccps.nic.in",
    "cms.nic.in",
    "niccdn.in",
    "/ccps/",
    "/nic-cms/",
    "nic.in/cms",
    "generator=ccps",
    "generator=nic cms",
    "x-powered-by=ccps",
    "x-cms=nic",
    "powered by nic cms",
    "powered by ccps",
    "national informatics centre cms",
    "<!-- ccps ",
    "<!-- nic cms",
]

CCPS_MEDIUM_CONFIDENCE = [
    "meghraj.nic",
    "nicnet",
    "niccloud",
    "goidirectory",
    "stqc certified",
]

CCPS_IGNORE = [
    "india.gov.in",
    "digitalindia.gov.in",
    "nic.in",
    "gov.in",
    "meity",
]

CCPS_HEADER_NAMES = [
    "x-powered-by",
    "x-generator",
    "server",
    "x-cms",
    "x-nic-cms",
]


# =============================================================================
# CONSTANTS — PER-01: Personalisation Consent (7.6.1) — MULTILINGUAL
# =============================================================================

PERSONALISATION_SELECTORS = [
    "[id*='personal']",   "[class*='personal']",
    "[id*='consent']",    "[class*='consent']",
    "[id*='preference']", "[class*='preference']",
    "[id*='customize']",  "[class*='customize']",
    "[id*='gdpr']",       "[class*='gdpr']",
    "[id*='tracking']",   "[class*='tracking']",
    "dialog", "[role='dialog']",
    ".modal", "#modal"
]

PERSONALISATION_KEYWORDS = [
    "personaliz", "personalise", "personalis",
    "tailor", "customise", "customize",
    "preference", "your experience",
    "remember your", "track your",
    "targeted", "recommendations for you"
]

# ── Multilingual Consent Choice Words ────────────────────────────────────────

CONSENT_CHOICE_WORDS_MULTILINGUAL = {
    # English
    "accept", "agree", "consent", "allow",
    "reject", "decline", "opt out", "opt-out",
    "i accept", "i agree", "i consent",
    "accept all", "reject all", "deny",
    "refuse", "do not accept",
    
    # Hindi (Devanagari script)
    "स्वीकार",      # accept (svīkār)
    "स्वीकृत",      # accepted (svīkṛt)
    "सहमत",         # agree (sahamat)
    "अनुमति",       # permission/allow (anumati)
    "अस्वीकार",     # reject (asvīkār)
    "मना",          # refuse (manā)
    "इनकार",        # denial (inkār)
    "सभी",          # all (sabhī)
    "कुकीज़",       # cookies (kukīz)
    "वैकल्पिक",     # optional (vaikalpik)
    "अनुकूलित",     # customize (anukūlit)
    
    # Transliterated Hindi (Latin script)
    "sweekar", "svikar", "sahmat",
    "anumati", "asveekar", "inkaar",
    "mana", "sabhi", "cookies",
    "vaikalpik", "anukoolit",
    
    # Common Indian English variations
    "ok", "okay", "got it", "understood",
    "no thanks", "not now", "later",
}

# ── Language Switcher / Translation Widget Selectors ──────────────────────────

LANG_SWITCHER_SELECTORS = [
    # Standard language selectors
    "[id*='lang']",      "[class*='lang']",
    "[id*='language']",  "[class*='language']",
    "select[name*='lang']",
    
    # URL-based language paths
    "a[href*='/hi/']",   "a[href*='/en/']",
    "a[href*='/ta/']",   "a[href*='/te/']",
    "a[href*='/bn/']",   "a[href*='/mr/']",
    "a[href*='/gu/']",   "a[href*='/kn/']",
    
    # Common CSS classes
    ".lang-switch", "#lang-switch",
    ".language-selector", "#language-selector",
    ".lang-dropdown", ".language-menu",
    
    # BHASHINI widget
    "[id*='bhashini']", "[class*='bhashini']",
    "[id*='translation']", "[class*='translation']",
    "[data-bhashini]",
    
    # Google Translate widget
    "#google_translate_element",
    ".goog-te-combo", ".goog-te-menu-value",
    
    # Microsoft Translator widget
    "[id*='MicrosoftTranslator']",
    "[class*='translator']",
    
    # Generic translation widgets
    "[role='combobox'][aria-label*='language']",
    "[role='combobox'][aria-label*='भाषा']",
]


# =============================================================================
# CONSTANTS — CK-03: Cookie Banner Compliance (7.6.1 / Req.42)
# =============================================================================

COOKIE_BANNER_SELECTORS = [
    "[id*='cookie']",        "[class*='cookie']",
    "[id*='consent']",       "[class*='consent']",
    "[id*='gdpr']",          "[class*='gdpr']",
    "[id*='cookiebar']",     "[class*='cookiebar']",
    "[id*='cookie-notice']", "[class*='cookie-notice']",
    "[id*='cookie-banner']", "[class*='cookie-banner']",
    ".cc-banner", ".cookie-law-info-bar",
    "#cookie-law-info-bar",  ".cookie-notice",
    "[aria-label*='cookie']",
    "[data-testid*='cookie']"
]

ACCEPT_PATTERNS = [
    "accept all", "accept cookies", "i accept",
    "allow all", "allow cookies", "agree",
    "accept", "ok", "got it", "allow"
]

REJECT_PATTERNS = [
    "reject all", "reject cookies", "i reject",
    "decline all", "decline cookies",
    "deny", "decline", "refuse",
    "do not accept", "reject", "no thanks"
]

CUSTOMIZE_PATTERNS = [
    "customize", "customise", "manage",
    "settings", "preferences", "options",
    "cookie settings", "manage cookies",
    "more options", "choose", "manage preferences"
]


# =============================================================================
# DRIVER HELPERS
# =============================================================================

def _find_chrome():
    candidates = [
        "google-chrome", "google-chrome-stable",
        "chromium", "chromium-browser"
    ]
    for c in candidates:
        path = shutil.which(c)
        if path:
            return path
    return None


def _find_chromedriver():
    candidates = [
        "chromedriver",
        "/usr/bin/chromedriver",
        "/usr/local/bin/chromedriver"
    ]
    for c in candidates:
        if c.startswith("/"):
            if os.path.exists(c):
                return c
        else:
            path = shutil.which(c)
            if path:
                return path
    return None


def create_driver():
    chrome_binary    = _find_chrome()
    chromedriver_bin = _find_chromedriver()

    if not chrome_binary:
        raise RuntimeError(
            "\nChrome/Chromium not found.\n"
            "Install: sudo apt install chromium"
        )

    if not chromedriver_bin:
        raise RuntimeError(
            "\nchromedriver not found.\n"
            "Install: sudo apt install chromium-driver"
        )

    opts = Options()
    opts.binary_location = chrome_binary
    opts.add_argument("--headless=new")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--disable-gpu")
    opts.add_argument("--window-size=1920,1080")
    opts.add_argument("--ignore-certificate-errors")
    opts.add_argument("--disable-blink-features=AutomationControlled")
    opts.add_argument(
        "--user-agent=Mozilla/5.0 "
        "(Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
    opts.add_experimental_option(
        "excludeSwitches", ["enable-automation"]
    )
    opts.add_experimental_option(
        "useAutomationExtension", False
    )

    service = Service(chromedriver_bin)
    driver  = webdriver.Chrome(service=service, options=opts)
    driver.set_page_load_timeout(60)

    driver.execute_cdp_cmd(
        "Page.addScriptToEvaluateOnNewDocument",
        {
            "source": """
                Object.defineProperty(navigator, 'webdriver',
                    { get: () => undefined });
                Object.defineProperty(navigator, 'plugins',
                    { get: () => [1, 2, 3] });
                Object.defineProperty(navigator, 'languages',
                    { get: () => ['en-IN', 'en', 'hi'] });
            """
        }
    )

    return driver


def wait_for_page(driver, timeout=30):
    WebDriverWait(driver, timeout).until(
        lambda d: d.execute_script(
            "return document.readyState"
        ) == "complete"
    )
    time.sleep(2)


# =============================================================================
# SHARED UTILITIES
# =============================================================================

def _page_text(driver):
    return driver.execute_script(
        "return document.body.innerText || '';"
    )


def _page_source(driver):
    return driver.page_source.lower()


def _is_visible(driver, element):
    try:
        return (
            element.is_displayed()
            and element.size["width"]  > 0
            and element.size["height"] > 0
        )
    except Exception:
        return False


def _first_visible(driver, selectors):
    for sel in selectors:
        try:
            for el in driver.find_elements(By.CSS_SELECTOR, sel):
                if _is_visible(driver, el):
                    return el
        except Exception:
            continue
    return None


def _try_parse_date(raw):
    raw = raw.strip()
    formats = [
        "%d %B %Y",  "%d %b %Y",
        "%d-%m-%Y",  "%d/%m/%Y",  "%d.%m.%Y",
        "%Y-%m-%d",  "%Y/%m/%d",
        "%B %d, %Y", "%b %d, %Y", "%d-%b-%Y",
    ]
    for fmt in formats:
        try:
            return datetime.strptime(raw, fmt)
        except ValueError:
            continue
    return None


def _match_button(candidates, patterns, driver):
    for el in candidates:
        try:
            text  = (el.text or "").lower().strip()
            aria  = (el.get_attribute("aria-label") or "").lower()
            value = (el.get_attribute("value")      or "").lower()
            combined = f"{text} {aria} {value}"
            for pat in patterns:
                if pat in combined and _is_visible(driver, el):
                    return el, text or aria or value
        except Exception:
            continue
    return None, None


# =============================================================================
# CL-06 HELPER FUNCTIONS
# =============================================================================

def strip_technical_tokens(text: str) -> str:
    for pattern in _EXCLUSION_PATTERNS:
        text = pattern.sub(" ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_human_readable_words(text: str) -> list:
    raw_words = re.findall(
        rf"\b[a-zA-Z]{{{SPELL_MIN_WORD_LEN},}}\b",
        text
    )
    cleaned = []
    for w in raw_words:
        if w.isupper():
            continue
        if re.search(r"\d", w):
            continue
        if "_" in w or "-" in w:
            continue
        if w.lower() in SPELL_WHITELIST:
            continue
        cleaned.append(w)
    return cleaned


# =============================================================================
# EXISTING CHECKS (CL-01 through CK-02)
# =============================================================================

def check_language_attribute(driver):
    lang = driver.execute_script(
        "return document.documentElement.lang || '';"
    ).strip().lower()
    if lang in ["en", "en-gb", "en-in"]:
        return {
            "passed": True,
            "reason": f"lang attribute correctly set to '{lang}'."
        }
    if not lang:
        return {
            "passed": False,
            "reason": "No lang attribute found on <html> element."
        }
    return {
        "passed": False,
        "reason": (
            f"Unexpected lang value: '{lang}'. "
            f"Expected en, en-GB, or en-IN."
        )
    }


def check_british_english(driver):
    text  = _page_text(driver).lower()
    found = [
        w for w in AMERICAN_SPELLINGS
        if re.search(rf"\b{w}\b", text)
    ]
    if found:
        return {
            "passed": False,
            "reason": (
                f"American spellings found: {', '.join(found[:8])}."
            )
        }
    return {
        "passed": True,
        "reason": "No common American spellings detected."
    }


def check_no_hinglish(driver):
    text  = _page_text(driver).lower()
    found = [
        w for w in HINGLISH_WORDS
        if re.search(rf"\b{w}\b", text)
    ]
    if found:
        return {
            "passed": False,
            "reason": (
                f"Hinglish words detected: {', '.join(found[:8])}."
            )
        }
    return {
        "passed": True,
        "reason": "No Hinglish words found."
    }


def check_https_links(driver):
    bad = [
        link.get_attribute("href")
        for link in driver.find_elements(By.TAG_NAME, "a")
        if (link.get_attribute("href") or "").startswith("http://")
    ]
    if bad:
        return {
            "passed": False,
            "reason": (
                f"{len(bad)} insecure HTTP link(s) found. "
                f"First: {bad[0]}"
            )
        }
    return {
        "passed": True,
        "reason": "All links use HTTPS."
    }


def check_language_switcher(driver):
    r = driver.execute_script("""
        var t = document.body.innerHTML.toLowerCase();
        return (
            t.includes('language') || t.includes('lang') ||
            t.includes('\u0939\u093f\u0928\u094d\u0926\u0940') ||
            t.includes('/hi/')     || t.includes('/en/')
        );
    """)
    if r:
        return {
            "passed": True,
            "reason": "Language switcher or multilingual navigation detected."
        }
    return {
        "passed": False,
        "reason": "No language switcher found."
    }


def check_hreflang(driver):
    tags = driver.find_elements(By.CSS_SELECTOR, "link[hreflang]")
    if tags:
        vals = [t.get_attribute("hreflang") for t in tags[:5]]
        return {
            "passed": True,
            "reason": f"hreflang tags found: {', '.join(vals)}."
        }
    found = driver.execute_script("""
        var found = [];
        document.querySelectorAll('link').forEach(function(l) {
            if (l.getAttribute('hreflang'))
                found.push(l.getAttribute('hreflang'));
        });
        return found;
    """)
    if found:
        return {
            "passed": True,
            "reason": (
                f"hreflang tags found via JS: {', '.join(found[:5])}."
            )
        }
    if "hreflang" in _page_source(driver):
        return {
            "passed": True,
            "reason": "hreflang attribute found in page source."
        }
    return {
        "passed": False,
        "reason": "No hreflang tags found."
    }


def check_cookie_banner(driver):
    if _first_visible(driver, COOKIE_BANNER_SELECTORS):
        return {
            "passed": True,
            "reason": "Cookie banner element detected in DOM."
        }
    text   = _page_text(driver).lower()
    source = _page_source(driver)
    for phrase in [
        "cookie", "accept cookies", "we use cookies",
        "this site uses cookies", "cookie consent"
    ]:
        if phrase in text or phrase in source:
            return {
                "passed": True,
                "reason": (
                    f"Cookie-related content detected ('{phrase}')."
                )
            }
    return {
        "passed": False,
        "reason": "No cookie banner found."
    }


def check_privacy_policy(driver):
    for link in driver.find_elements(By.TAG_NAME, "a"):
        text = (link.text or "").lower()
        href = (link.get_attribute("href") or "").lower()
        if "privacy" in text or "privacy" in href:
            return {
                "passed": True,
                "reason": (
                    f"Privacy policy link found: "
                    f"'{link.text.strip()}'."
                )
            }
    if "privacy" in _page_source(driver):
        return {
            "passed": True,
            "reason": "Privacy policy reference found in page source."
        }
    return {
        "passed": False,
        "reason": "No privacy policy link found."
    }


# =============================================================================
# CHECK — CL-05: Content Completeness & Freshness (A.5.6)
# =============================================================================

def check_content_freshness(driver, base_url=None):
    issues = []
    info   = []
    page_text = _page_text(driver)
    source    = driver.page_source

    meta_values = driver.execute_script("""
        var names = [
            'date','last-modified','revised',
            'dcterms.modified','article:modified_time'
        ];
        var found = [];
        document.querySelectorAll('meta').forEach(function(m) {
            var n = (m.getAttribute('name')     || '').toLowerCase();
            var p = (m.getAttribute('property') || '').toLowerCase();
            var c =  m.getAttribute('content')  || '';
            if (c && (names.indexOf(n) > -1 || names.indexOf(p) > -1))
                found.push(c);
        });
        return found;
    """)

    date_found = None
    for mv in (meta_values or []):
        d = _try_parse_date(mv[:10])
        if d:
            date_found = d
            break

    if not date_found:
        combined = page_text + " " + source
        for pat in LAST_UPDATED_PATTERNS:
            for match in re.findall(pat, combined, re.IGNORECASE):
                d = _try_parse_date(match)
                if d:
                    date_found = d
                    break
            if date_found:
                break

    if date_found:
        age_months = (datetime.now() - date_found).days / 30.44
        age_label  = date_found.strftime("%d %b %Y")
        if age_months > STALE_THRESHOLD_MONTHS:
            issues.append(
                f"Last updated date ({age_label}) is "
                f"{int(age_months)} months old — "
                f"exceeds {STALE_THRESHOLD_MONTHS}-month threshold."
            )
        else:
            info.append(
                f"Last updated: {age_label} "
                f"({int(age_months)} months ago) — current."
            )
    else:
        issues.append(
            "No 'Last Updated' date detected. "
            "DBIM A.5.6 requires a visible publication date."
        )

    text_lower       = page_text.lower()
    placeholder_hits = []
    for pat in PLACEHOLDER_PATTERNS:
        if re.search(pat, text_lower, re.IGNORECASE):
            readable = re.sub(
                r"\\[bsS+?]", "", pat
            ).replace(r"\b", "").strip()
            placeholder_hits.append(readable)
    if placeholder_hits:
        issues.append(
            f"Placeholder / incomplete content detected: "
            f"{', '.join(placeholder_hits[:5])}."
        )

    empty_data = driver.execute_script(
        """
        var selectors = arguments[0];
        var count = 0;
        var details = [];
        selectors.forEach(function(sel) {
            document.querySelectorAll(sel).forEach(function(el, i) {
                var txt   = (el.innerText || '').trim();
                var hasImg = el.querySelectorAll('img').length > 0;
                var hasAV  = el.querySelectorAll(
                    'video,audio,iframe,embed'
                ).length > 0;
                if (txt.length < 10 && !hasImg && !hasAV) {
                    count++;
                    details.push(sel + '[' + i + ']');
                }
            });
        });
        return { count: count, details: details.slice(0, 5) };
        """,
        CONTENT_SECTION_SELECTORS
    )
    empty_count = (empty_data or {}).get("count", 0)
    if empty_count > 0:
        details = (empty_data or {}).get("details", [])
        issues.append(
            f"{empty_count} empty content section(s) detected "
            f"({', '.join(details[:3])})."
        )

    if base_url:
        base_netloc = urlparse(base_url).netloc
        hrefs       = set()
        for a in driver.find_elements(By.TAG_NAME, "a"):
            href = a.get_attribute("href") or ""
            if not href:
                continue
            if href.startswith(
                ("javascript", "mailto", "tel", "#")
            ):
                continue
            parsed = urlparse(href)
            if (
                not parsed.netloc
                or parsed.netloc == base_netloc
            ):
                hrefs.add(href)
        broken  = []
        session = requests.Session()
        session.verify = False
        for href in list(hrefs)[:MAX_INTERNAL_LINKS_TO_CHECK]:
            try:
                full = urljoin(base_url, href)
                resp = session.head(
                    full, timeout=8, allow_redirects=True
                )
                if resp.status_code >= 400:
                    broken.append(
                        f"{full} [HTTP {resp.status_code}]"
                    )
            except Exception:
                pass
        if broken:
            issues.append(
                f"{len(broken)} broken internal link(s): "
                f"{'; '.join(broken[:3])}."
            )
        elif hrefs:
            checked = min(len(hrefs), MAX_INTERNAL_LINKS_TO_CHECK)
            info.append(
                f"{checked} internal link(s) checked — none broken."
            )

    if issues:
        return {"passed": False, "reason": " | ".join(issues)}
    return {
        "passed": True,
        "reason": " | ".join(info) if info
                  else "Content is complete and current."
    }


# =============================================================================
# CHECK — CL-06: Grammar & Spelling Validation (7.1.3.3)
# =============================================================================

def check_grammar_spelling(driver):
    issues = []
    info   = []
    raw_page_text = _page_text(driver)
    if not raw_page_text or len(raw_page_text.strip()) < 50:
        return {
            "passed": False,
            "reason": "Page text too short or empty to validate."
        }
    clean_text = strip_technical_tokens(raw_page_text)
    if len(clean_text.strip()) < 20:
        return {
            "passed": True,
            "reason": (
                "After removing technical tokens (emails, URLs, "
                "identifiers), insufficient plain text remains to "
                "spell-check. No human-readable errors detected."
            )
        }

    if _SPELL_OK:
        spell = SpellChecker(language="en")
        words_to_check = extract_human_readable_words(clean_text)
        unique_words = list(set(w.lower() for w in words_to_check))
        misspelled_raw = spell.unknown(unique_words)
        title_cased_lowers = {
            w.lower()
            for w in words_to_check
            if w and w[0].isupper()
        }
        real_errors = []
        for misspelled_word in misspelled_raw:
            if misspelled_word in title_cased_lowers:
                continue
            if misspelled_word in SPELL_WHITELIST:
                continue
            if len(misspelled_word) < SPELL_MIN_WORD_LEN:
                continue
            real_errors.append(misspelled_word)
        real_errors = real_errors[:SPELL_MAX_SAMPLE]
        if real_errors:
            issues.append(
                f"{len(real_errors)} spelling error(s) detected in "
                f"visible content. "
                f"Sample: {', '.join(real_errors[:8])}."
            )
        else:
            info.append(
                "No spelling errors detected in visible content "
                "(technical tokens excluded from check)."
            )
    else:
        text_lower   = clean_text.lower()
        found_errors = [
            f"'{wrong}' → '{right}'"
            for wrong, right in FALLBACK_MISSPELLINGS
            if re.search(rf"\b{wrong}\b", text_lower)
        ]
        if found_errors:
            issues.append(
                f"Spelling errors found in visible content: "
                f"{', '.join(found_errors)}."
            )
        else:
            info.append(
                "Basic spell check passed on visible content "
                "(install pyspellchecker for full coverage)."
            )

    if _GRAMMAR_OK:
        try:
            tool = language_tool_python.LanguageTool("en-GB")
            grammar_sample = clean_text[:3000]
            matches = tool.check(grammar_sample)
            tool.close()
            noise_rules = {
                "WHITESPACE_RULE",
                "EN_QUOTES",
                "COMMA_PARENTHESIS_WHITESPACE",
                "UPPERCASE_SENTENCE_START",
                "WORD_CONTAINS_UNDERSCORE",
                "EN_UNPAIRED_BRACKETS",
                "DASH_RULE",
                "MULTIPLICATION_SIGN",
                "UNIT_SPACE",
            }
            grammar_errors = [
                m for m in matches
                if m.ruleId not in noise_rules
            ]
            if grammar_errors:
                sample_msgs = list(
                    {m.message for m in grammar_errors[:5]}
                )
                issues.append(
                    f"{len(grammar_errors)} grammar issue(s) detected "
                    f"in visible content. "
                    f"Sample: {'; '.join(sample_msgs[:3])}."
                )
            else:
                info.append(
                    "No grammar issues detected in visible content."
                )
        except Exception as e:
            info.append(
                f"Grammar check skipped (LanguageTool error: {e})."
            )
    else:
        info.append(
            "Grammar check skipped "
            "(install language_tool_python for full coverage)."
        )

    words    = clean_text.lower().split()
    repeated = []
    for i in range(len(words) - 1):
        a = re.sub(r"[^a-z]", "", words[i])
        b = re.sub(r"[^a-z]", "", words[i + 1])
        if a and a == b and len(a) > 2:
            repeated.append(a)
    if repeated:
        issues.append(
            f"Repeated consecutive word(s) in visible content: "
            f"{', '.join(list(set(repeated))[:5])}."
        )

    all_caps = re.findall(
        rf"\b[A-Z]{{{EXCESSIVE_CAP_MIN_LEN},}}\b",
        raw_page_text
    )
    excessive = [
        w for w in all_caps
        if w not in KNOWN_CAPS_TOKENS
    ]
    if len(excessive) > 3:
        issues.append(
            f"{len(excessive)} word(s) with excessive capitalisation "
            f"detected: {', '.join(excessive[:5])}."
        )

    sentences = re.split(r"[.!?]+", clean_text)
    malformed = [
        s.strip()[:60]
        for s in sentences
        if 1 < len(s.strip().split()) < 3
        and len(s.strip()) > 5
    ]
    if len(malformed) > 5:
        issues.append(
            f"{len(malformed)} potentially malformed / fragment "
            f"sentence(s) detected in visible content."
        )

    if issues:
        return {
            "passed": False,
            "reason": " | ".join(issues)
        }
    return {
        "passed": True,
        "reason": (
            " | ".join(info)
            if info
            else "Grammar and spelling checks passed."
        )
    }


# =============================================================================
# CHECK — CCPS-01: CCPS Integration (7.4)
# =============================================================================

def check_ccps_integration(driver):
    high_confidence_evidence   = []
    medium_confidence_evidence = []
    
    current_url    = driver.current_url
    current_domain = urlparse(current_url).netloc.lower()

    def is_internal_url(url):
        if not url or url.startswith(("javascript:", "mailto:", "tel:", "#")):
            return False
        try:
            parsed = urlparse(url)
            if not parsed.netloc:
                return True
            return parsed.netloc.lower() == current_domain
        except Exception:
            return False

    def add_evidence(level, source_type, marker, location, context=""):
        display = f"[{source_type}] {marker}"
        if location:
            display += f" → {location[:80]}"
        if context:
            display += f" | ...{context[:60]}..."
        
        if level == "HIGH":
            high_confidence_evidence.append(display)
        else:
            medium_confidence_evidence.append(display)

    for script in driver.find_elements(By.TAG_NAME, "script"):
        src = (script.get_attribute("src") or "").strip()
        
        if src:
            src_lower = src.lower()
            
            for marker in CCPS_HIGH_CONFIDENCE:
                if marker in src_lower:
                    add_evidence(
                        "HIGH",
                        "Script src",
                        marker,
                        src
                    )
                    break
            
            if is_internal_url(src) or "nic.in" in urlparse(src).netloc:
                for marker in CCPS_MEDIUM_CONFIDENCE:
                    if marker in src_lower:
                        add_evidence(
                            "MEDIUM",
                            "Script src",
                            marker,
                            src
                        )
                        break
        
        try:
            inline = (script.get_attribute("innerHTML") or "")[:1000]
            inline_lower = inline.lower()
            
            for marker in CCPS_HIGH_CONFIDENCE:
                if marker in inline_lower:
                    pos = inline_lower.find(marker)
                    start = max(0, pos - 30)
                    end = min(len(inline), pos + len(marker) + 30)
                    context = inline[start:end].strip()
                    
                    add_evidence(
                        "HIGH",
                        "Inline script",
                        marker,
                        "(embedded in page)",
                        context
                    )
                    break
        except Exception:
            pass

    for meta in driver.find_elements(By.TAG_NAME, "meta"):
        name    = (meta.get_attribute("name")     or "").lower()
        prop    = (meta.get_attribute("property") or "").lower()
        content = (meta.get_attribute("content")  or "").lower()
        
        combined = f"{name} {prop} {content}"
        
        for marker in CCPS_HIGH_CONFIDENCE:
            if marker in combined:
                add_evidence(
                    "HIGH",
                    "Meta tag",
                    marker,
                    f"name='{name}' content='{content[:50]}'"
                )
                break

    for tag, attr in [
        ("link",   "href"),
        ("iframe", "src"),
        ("form",   "action"),
    ]:
        for el in driver.find_elements(By.TAG_NAME, tag)[:100]:
            url = (el.get_attribute(attr) or "").strip()
            
            if not url or not is_internal_url(url):
                continue
            
            url_lower = url.lower()
            
            for marker in CCPS_HIGH_CONFIDENCE:
                if marker in url_lower:
                    add_evidence(
                        "HIGH",
                        f"{tag.upper()} {attr}",
                        marker,
                        url
                    )
                    break
            
            for marker in CCPS_MEDIUM_CONFIDENCE:
                if marker in url_lower:
                    add_evidence(
                        "MEDIUM",
                        f"{tag.upper()} {attr}",
                        marker,
                        url
                    )
                    break

    try:
        full_html = driver.execute_script(
            "return document.documentElement.outerHTML;"
        ).lower()
        
        for marker in CCPS_HIGH_CONFIDENCE:
            if marker in full_html:
                pos = full_html.find(marker)
                start = max(0, pos - 40)
                end = min(len(full_html), pos + len(marker) + 40)
                context = full_html[start:end].strip()
                
                add_evidence(
                    "HIGH",
                    "HTML source",
                    marker,
                    "(embedded in page HTML)",
                    context
                )
        
        for marker in CCPS_MEDIUM_CONFIDENCE:
            if marker in full_html:
                pos = full_html.find(marker)
                start = max(0, pos - 40)
                end = min(len(full_html), pos + len(marker) + 40)
                context = full_html[start:end].strip()
                
                add_evidence(
                    "MEDIUM",
                    "HTML source",
                    marker,
                    "(embedded in page HTML)",
                    context
                )
                
    except Exception:
        try:
            source_sample = driver.page_source.lower()
            
            for marker in CCPS_HIGH_CONFIDENCE:
                if marker in source_sample:
                    add_evidence(
                        "HIGH",
                        "HTML source (fallback)",
                        marker,
                        "(page source)"
                    )
        except Exception:
            pass

    try:
        resp = requests.head(
            current_url,
            timeout=8,
            verify=False,
            allow_redirects=True
        )
        
        for hdr in CCPS_HEADER_NAMES:
            val = resp.headers.get(hdr, "").lower()
            if not val:
                continue
            
            for marker in CCPS_HIGH_CONFIDENCE:
                if marker in val:
                    add_evidence(
                        "HIGH",
                        f"HTTP header [{hdr}]",
                        marker,
                        val[:60]
                    )
                    break
                    
    except Exception:
        pass

    high_count   = len(high_confidence_evidence)
    medium_count = len(medium_confidence_evidence)
    
    high_unique   = list(dict.fromkeys(high_confidence_evidence))
    medium_unique = list(dict.fromkeys(medium_confidence_evidence))

    if high_count >= 1:
        return {
            "passed": True,
            "reason": (
                f"CCPS integration CONFIRMED — "
                f"{high_count} HIGH-confidence indicator(s) found: "
                f"{' | '.join(high_unique[:3])}"
                + (f" (+{medium_count} medium-confidence)" if medium_count > 0 else "")
            )
        }
    
    if medium_count >= 3:
        return {
            "passed": True,
            "reason": (
                f"CCPS integration LIKELY — "
                f"{medium_count} MEDIUM-confidence indicator(s) found: "
                f"{' | '.join(medium_unique[:5])}. "
                "No high-confidence markers detected, but multiple "
                "NIC infrastructure references suggest CCPS hosting."
            )
        }

    if 1 <= medium_count <= 2:
        return {
            "passed": False,
            "reason": (
                f"CCPS integration INCONCLUSIVE — "
                f"{medium_count} MEDIUM-confidence indicator(s) found: "
                f"{' | '.join(medium_unique)}. "
                "Insufficient evidence to confirm CCPS integration. "
                "Manual verification required: inspect page source for "
                "cms.nic.in, ccps.nic.in, or 'Powered by NIC CMS'. "
                "DBIM 7.4 requires Central Content Publishing System integration."
            )
        }
    
    return {
        "passed": False,
        "reason": (
            "No CCPS integration evidence found. "
            "No high-confidence or medium-confidence indicators detected. "
            "Expected indicators: cms.nic.in, ccps.nic.in, /ccps/ paths, "
            "'Powered by NIC CMS', or X-Powered-By: CCPS header. "
            "DBIM 7.4 requires Central Content Publishing System integration. "
            "Manual verification: check page source (Ctrl+U) and Network tab "
            "for CMS-related resources."
        )
    }


# =============================================================================
# CHECK — PER-01: Personalisation Consent (7.6.1) — MULTILINGUAL
# =============================================================================

def check_personalisation_consent(driver):
    """
    PER-01 — Clear consent for personalisation must be obtained
    in the user's preferred language (DBIM 7.6.1).

    MULTILINGUAL DETECTION
    ──────────────────────
    Detects consent controls in:
      • English (Accept, Reject, Customize)
      • Hindi Devanagari (स्वीकार, अस्वीकार, अनुकूलित)
      • Transliterated Hindi (sweekar, asveekar, anukoolit)
      • Other Indian languages via Unicode range detection
    
    LANGUAGE SELECTOR DETECTION
    ───────────────────────────
    Recognizes:
      • Standard language dropdowns (<select>, href="/hi/", etc.)
      • BHASHINI translation widgets (Government of India)
      • Google Translate / Microsoft Translator widgets
      • Custom language switchers
    
    PASS CRITERIA
    ─────────────
    1. Explicit consent choice controls exist (any language), AND
    2. Language selection mechanism is available
    
    Sub-tests
    ─────────
      1. Detect personalisation consent element / keywords.
      2. Verify explicit consent choice wording (multilingual).
      3. Detect language switcher / translation widget.
      4. Optional: Test language switch behavior (if possible).
    """

    issues = []
    info   = []
    
    page_text = _page_text(driver)
    page_text_lower = page_text.lower()
    
    full_html = ""
    try:
        full_html = driver.execute_script(
            "return document.documentElement.outerHTML;"
        ).lower()
    except Exception:
        full_html = driver.page_source.lower()

    consent_el = _first_visible(driver, PERSONALISATION_SELECTORS)

    personalisation_found = any(
        kw in page_text_lower
        for kw in PERSONALISATION_KEYWORDS
    )

    if not consent_el and not personalisation_found:
        return {
            "passed": True,
            "reason": (
                "No personalisation mechanism detected. "
                "If the site does not personalise content, "
                "this requirement may not apply. "
                "Verify manually with site owner."
            )
        }

    if personalisation_found:
        info.append("Personalisation-related content detected.")

    consent_found_in_language = None
    
    english_consent = any(
        kw in page_text_lower
        for kw in ["accept", "reject", "decline", "consent", "agree"]
    )
    
    hindi_consent = any(
        kw in page_text
        for kw in ["स्वीकार", "अस्वीकार", "सहमत", "अनुमति", "मना"]
    )
    
    transliterated_consent = any(
        kw in page_text_lower
        for kw in ["sweekar", "asveekar", "sahmat", "anumati"]
    )
    
    multilingual_consent = any(
        kw in page_text_lower or kw in page_text
        for kw in CONSENT_CHOICE_WORDS_MULTILINGUAL
    )

    if multilingual_consent:
        if hindi_consent:
            consent_found_in_language = "Hindi (Devanagari)"
        elif transliterated_consent:
            consent_found_in_language = "Hindi (transliterated)"
        elif english_consent:
            consent_found_in_language = "English"
        else:
            consent_found_in_language = "Detected (language not identified)"
        
        info.append(
            f"Explicit consent choice wording found in {consent_found_in_language}."
        )
    else:
        issues.append(
            "No explicit consent choice wording found in any supported language "
            "(English, Hindi, or transliterated variants). "
            "Expected: Accept/Reject/Customize or स्वीकार/अस्वीकार/अनुकूलित."
        )

    lang_switcher = _first_visible(driver, LANG_SWITCHER_SELECTORS)
    
    bhashini_detected = False
    google_translate_detected = False
    other_translator_detected = False
    
    if (
        "bhashini" in full_html
        or "bhashini" in page_text_lower
        or driver.find_elements(By.CSS_SELECTOR, "[data-bhashini]")
    ):
        bhashini_detected = True
        info.append("BHASHINI translation widget detected.")
    
    if (
        "google_translate_element" in full_html
        or "googletrans" in full_html
        or driver.find_elements(By.CSS_SELECTOR, ".goog-te-combo")
    ):
        google_translate_detected = True
        info.append("Google Translate widget detected.")
    
    if "microsofttranslator" in full_html:
        other_translator_detected = True
        info.append("Microsoft Translator widget detected.")
    
    url_based_lang = False
    try:
        current_url = driver.current_url.lower()
        if any(lang_path in current_url for lang_path in ["/hi/", "/en/", "/ta/", "/bn/"]):
            url_based_lang = True
            info.append("URL-based language selection detected (e.g., /hi/, /en/).")
    except Exception:
        pass

    language_mechanism_found = (
        lang_switcher is not None
        or bhashini_detected
        or google_translate_detected
        or other_translator_detected
        or url_based_lang
    )

    if not language_mechanism_found:
        issues.append(
            "No language selection mechanism found. "
            "Expected: language dropdown, BHASHINI widget, Google Translate, "
            "or URL-based language paths (/hi/, /en/). "
            "DBIM 7.6.1 requires consent in the user's preferred language."
        )

    if lang_switcher:
        baseline = page_text[:2000]
        adaptive = False
        
        try:
            lang_switcher.click()
            time.sleep(2)
            
            post_text = _page_text(driver)
            
            script_ranges = {
                "Devanagari": r"[\u0900-\u097F]",
                "Tamil": r"[\u0B80-\u0BFF]",
                "Telugu": r"[\u0C00-\u0C7F]",
                "Bengali": r"[\u0980-\u09FF]",
                "Gujarati": r"[\u0A80-\u0AFF]",
                "Kannada": r"[\u0C80-\u0CFF]",
            }
            
            detected_script = None
            for script_name, pattern in script_ranges.items():
                if re.search(pattern, post_text):
                    detected_script = script_name
                    break
            
            if detected_script:
                adaptive = True
                info.append(
                    f"Consent text adapts after language switch — "
                    f"{detected_script} script detected."
                )
            elif post_text[:2000] != baseline:
                adaptive = True
                info.append(
                    "Page content changed after language switch."
                )
            
            driver.back()
            wait_for_page(driver)
            
        except Exception as e:
            info.append(
                f"Language switch test incomplete: {e}."
            )
        
        if not adaptive:
            issues.append(
                "Language switcher present but consent text did not "
                "adapt to the selected language. "
                "DBIM 7.6.1 requires consent in the user's preferred language."
            )

    if issues:
        return {
            "passed": False,
            "reason": " | ".join(issues)
        }

    return {
        "passed": True,
        "reason": (
            " | ".join(info)
            if info
            else "Personalisation consent with multilingual support detected."
        )
    }


# =============================================================================
# CHECK — CK-03: Cookie Banner Compliance
# =============================================================================

def check_cookie_banner_compliance(driver):
    issues = []
    info   = []
    banner = None
    for sel in COOKIE_BANNER_SELECTORS:
        try:
            for el in driver.find_elements(By.CSS_SELECTOR, sel):
                if _is_visible(driver, el):
                    banner = el
                    break
        except Exception:
            continue
        if banner:
            break
    if not banner:
        return {
            "passed": False,
            "reason": (
                "No cookie banner element detected. "
                "DBIM Req.42 requires a cookie consent banner at "
                "the bottom of the page with Accept, Reject, "
                "and Customize options."
            )
        }
    info.append("Cookie banner element found.")
    try:
        pos = driver.execute_script(
            """
            var el   = arguments[0];
            var st   = window.getComputedStyle(el);
            var rect = el.getBoundingClientRect();
            var vh   = window.innerHeight;
            return {
                position : st.position,
                bottom   : st.bottom,
                rectTop  : rect.top,
                vpHeight : vh,
                isFixed  : st.position === 'fixed',
                isSticky : st.position === 'sticky'
            };
            """,
            banner
        )
        is_fixed_or_sticky = (
            pos.get("isFixed") or pos.get("isSticky")
        )
        rect_top  = pos.get("rectTop",  0)
        vp_height = pos.get("vpHeight", 768)
        in_bottom_zone = rect_top > (vp_height * 0.65)
        css_bottom     = (pos.get("bottom") or "").strip()
        at_css_bottom  = css_bottom in ("0px", "0")
        positioned_correctly = (
            (is_fixed_or_sticky and at_css_bottom)
            or in_bottom_zone
        )
        if positioned_correctly:
            info.append(
                f"Banner position: {pos.get('position')} — "
                f"correctly anchored at bottom."
            )
        else:
            issues.append(
                f"Cookie banner is NOT fixed at the bottom "
                f"(CSS position: {pos.get('position')}, "
                f"CSS bottom: {css_bottom}, "
                f"viewport top offset: {int(rect_top)}px / "
                f"{vp_height}px). "
                f"DBIM Req.42 requires the banner at page bottom."
            )
    except Exception as e:
        issues.append(f"Position check could not be completed: {e}.")
    try:
        banner_btns = banner.find_elements(
            By.CSS_SELECTOR,
            "button, a, input[type='button'], "
            "input[type='submit'], [role='button'], [onclick]"
        )
    except Exception:
        banner_btns = []
    try:
        all_btns = driver.find_elements(
            By.CSS_SELECTOR,
            "button, a[role='button'], input[type='button']"
        )
    except Exception:
        all_btns = []
    banner_set  = set(banner_btns)
    extra_btns  = [b for b in all_btns if b not in banner_set]
    candidates  = banner_btns + extra_btns
    accept_el, accept_text = _match_button(
        candidates, ACCEPT_PATTERNS, driver
    )
    if accept_el:
        info.append(f"Accept button found: '{accept_text}'.")
    else:
        issues.append(
            "No visible Accept / Allow button found. "
            "DBIM Req.42 requires an explicit 'Accept' option."
        )
    reject_el, reject_text = _match_button(
        candidates, REJECT_PATTERNS, driver
    )
    if reject_el:
        info.append(f"Reject button found: '{reject_text}'.")
    else:
        issues.append(
            "No visible Reject / Decline button found. "
            "DBIM Req.42 requires an explicit 'Reject' option."
        )
    customize_el, customize_text = _match_button(
        candidates, CUSTOMIZE_PATTERNS, driver
    )
    if customize_el:
        info.append(f"Customize button found: '{customize_text}'.")
    else:
        issues.append(
            "No visible Customize / Settings button found. "
            "DBIM Req.42 requires a 'Customize' or "
            "'Manage Preferences' option."
        )
    for label, el in [
        ("Accept",    accept_el),
        ("Reject",    reject_el),
        ("Customize", customize_el),
    ]:
        if el:
            try:
                if not el.is_enabled():
                    issues.append(
                        f"{label} button is disabled."
                    )
                elif not el.is_displayed():
                    issues.append(
                        f"{label} button is not displayed."
                    )
            except Exception as e:
                issues.append(
                    f"{label} button check failed: {e}."
                )
    if issues:
        return {"passed": False, "reason": " | ".join(issues)}
    return {
        "passed": True,
        "reason": " | ".join(info) if info
                  else (
                      "Cookie banner with Accept, Reject, and Customize "
                      "buttons detected at page bottom."
                  )
    }


# =============================================================================
# RESULT CLASS
# =============================================================================

class AuditResult:
    def __init__(
        self, tc_id, category, name,
        description, result, reason, screenshot=None
    ):
        self.tc_id       = tc_id
        self.category    = category
        self.name        = name
        self.description = description
        self.result      = result
        self.reason      = reason
        self.screenshot  = screenshot


# =============================================================================
# AUDIT ENGINE
# =============================================================================

def run_audit(url, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    ss_dir = os.path.join(output_dir, "screenshots")
    os.makedirs(ss_dir, exist_ok=True)
    driver  = create_driver()
    results = []

    def snap(name):
        path = os.path.join(ss_dir, name)
        try:
            driver.save_screenshot(path)
            return path
        except Exception:
            return None

    def add(tc_id, category, name, description, fn):
        try:
            r = fn()
            results.append(AuditResult(
                tc_id, category, name, description,
                "PASS" if r["passed"] else "FAIL",
                r["reason"],
                snap(f"{tc_id}.png")
            ))
        except Exception as e:
            results.append(AuditResult(
                tc_id, category, name, description,
                "SKIP",
                f"Check skipped — error: {e}",
                snap(f"{tc_id}_error.png")
            ))

    try:
        print(f"\n[*] Opening {url}")
        driver.get(url)
        wait_for_page(driver)

        print("[1] Content & Language checks")
        add("CL-01", "Content & Language",
            "Language Attribute",
            "HTML lang attribute (DBIM 7.1.3.1)",
            lambda: check_language_attribute(driver))
        add("CL-02", "Content & Language",
            "British English",
            "No American spellings (DBIM 7.1.3.2)",
            lambda: check_british_english(driver))
        add("CL-03", "Content & Language",
            "No Hinglish",
            "Content must not contain Hinglish (DBIM 7.1.3.3)",
            lambda: check_no_hinglish(driver))
        add("CL-04", "Content & Language",
            "HTTPS Links",
            "All hyperlinks must use HTTPS",
            lambda: check_https_links(driver))
        add("CL-05", "Content & Language",
            "Content Completeness and Freshness",
            "All content must be complete and up to date (DBIM A.5.6)",
            lambda: check_content_freshness(driver, base_url=url))
        add("CL-06", "Content & Language",
            "Grammar and Spelling Validation",
            "Free from spelling / grammar errors; no Hinglish "
            "(DBIM 7.1.3.3)",
            lambda: check_grammar_spelling(driver))

        print("[2] Multilingual checks")
        add("ML-01", "Multilingual",
            "Language Switcher",
            "Language switcher must be visible (DBIM 7.2)",
            lambda: check_language_switcher(driver))
        add("ML-02", "Multilingual",
            "hreflang Tags",
            "hreflang implementation (DBIM 7.2)",
            lambda: check_hreflang(driver))

        print("[3] CCPS Integration checks")
        add("CCPS-01", "CCPS Integration",
            "Central Content Publishing System Integration",
            "Website must be integrated with NIC CCPS (DBIM 7.4)",
            lambda: check_ccps_integration(driver))

        print("[4] Cookie & Personalisation checks")
        add("CK-01", "Cookie & Personalisation",
            "Cookie Banner Existence",
            "Cookie banner must be present (DBIM 7.6.1)",
            lambda: check_cookie_banner(driver))
        add("CK-02", "Cookie & Personalisation",
            "Privacy Policy Link",
            "Privacy policy link must be present (DBIM 7.6)",
            lambda: check_privacy_policy(driver))
        add("CK-03", "Cookie & Personalisation",
            "Cookie Banner Full Compliance",
            "Banner at bottom with Accept / Reject / Customize "
            "(DBIM Req.42 / 7.6.1)",
            lambda: check_cookie_banner_compliance(driver))
        add("PER-01", "Cookie & Personalisation",
            "Personalisation Consent Language Validation",
            "Consent for personalisation in user's preferred language "
            "(DBIM 7.6.1)",
            lambda: check_personalisation_consent(driver))

        pass_c = sum(1 for r in results if r.result == "PASS")
        fail_c = sum(1 for r in results if r.result == "FAIL")
        skip_c = sum(1 for r in results if r.result == "SKIP")
        results.append(AuditResult(
            "SUMMARY", "Summary",
            "Overall Result",
            "Aggregate DBIM compliance result",
            "PASS" if fail_c == 0 else "FAIL",
            (
                f"{pass_c} passed, "
                f"{fail_c} failed, "
                f"{skip_c} skipped — "
                f"{pass_c + fail_c + skip_c} total checks."
            )
        ))
    finally:
        driver.quit()
    return results


# =============================================================================
# REPORT GENERATION
# =============================================================================

def shade_cell(cell, color):
    tc  = cell._element.get_or_add_tcPr()
    shd = tc.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color, qn("w:val"): "clear"}
    )
    tc.append(shd)


def generate_report(url, results, output_dir):
    doc   = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    t = doc.add_heading("DBIM Website Compliance Report", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"URL: {url}\n").bold = True
    meta.add_run(
        f"Generated: "
        f"{datetime.now().strftime('%d %B %Y  %H:%M:%S')}\n"
    )
    meta.add_run(
        "Standard: Digital Brand Identity Manual (DBIM)\n"
    )
    doc.add_page_break()
    colors = {
        "PASS": "C6EFCE",
        "FAIL": "FFC7CE",
        "SKIP": "FFF2CC"
    }
    categories = list(
        dict.fromkeys(r.category for r in results)
    )
    for cat in categories:
        doc.add_heading(cat, level=1)
        for res in (r for r in results if r.category == cat):
            doc.add_heading(
                f"{res.tc_id} — {res.name}", level=2
            )
            tbl = doc.add_table(rows=4, cols=2)
            tbl.style = "Table Grid"
            fields = [
                ("Test ID",     res.tc_id),
                ("Description", res.description),
                ("Result",      res.result),
                ("Reason",      res.reason),
            ]
            for i, (k, v) in enumerate(fields):
                tbl.rows[i].cells[0].text = k
                tbl.rows[i].cells[1].text = str(v)
                for para in tbl.rows[i].cells[0].paragraphs:
                    for run in para.runs:
                        run.bold = True
                if k == "Result":
                    shade_cell(
                        tbl.rows[i].cells[1],
                        colors.get(v, "FFFFFF")
                    )
            if (
                res.screenshot
                and os.path.exists(res.screenshot)
            ):
                doc.add_paragraph()
                cap = doc.add_paragraph(
                    f"Screenshot Evidence — "
                    f"{res.tc_id}: {res.name}"
                )
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_picture(
                    res.screenshot, width=Inches(5.5)
                )
                doc.paragraphs[-1].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )
            doc.add_paragraph()
    parsed    = urlparse(url)
    safe_name = re.sub(r"[^a-zA-Z0-9]", "_", parsed.netloc)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path      = os.path.join(
        output_dir,
        f"DBIM_Report_{safe_name}_{timestamp}.docx"
    )
    doc.save(path)
    return path


# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="DBIM Website Compliance Checker"
    )
    parser.add_argument(
        "-o", "--output",
        default="dbim_output",
        help="Output directory (default: dbim_output)"
    )
    args = parser.parse_args()
    url = input("Enter website URL: ").strip()
    if not url.startswith("http"):
        url = "https://" + url
    print("=" * 60)
    print("DBIM WEBSITE COMPLIANCE CHECKER")
    print("=" * 60)
    print(f"Target : {url}")
    print(f"Output : {args.output}")
    print("=" * 60)
    results = run_audit(url, args.output)
    print("\n[*] Generating report...")
    report = generate_report(url, results, args.output)
    print(f"[+] Report saved: {report}")
    print("\nRESULTS")
    print("-" * 70)
    icons = {"PASS": "✅", "FAIL": "❌", "SKIP": "⏭️"}
    for r in results:
        print(
            f"{icons.get(r.result, '?')} "
            f"{r.tc_id:<12s} "
            f"{r.result:<6s} "
            f"{r.name}"
        )
    print("-" * 70)
    fail_c = sum(1 for r in results if r.result == "FAIL")
    skip_c = sum(1 for r in results if r.result == "SKIP")
    pass_c = sum(1 for r in results if r.result == "PASS")
    print(
        f"\nTotal: {pass_c} passed | "
        f"{fail_c} failed | "
        f"{skip_c} skipped"
    )
    if fail_c:
        print(f"\n[!] {fail_c} compliance check(s) FAILED.")
    else:
        print("\n[✓] All checks passed.")
    return 0 if fail_c == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())