#!/usr/bin/env python3
"""
====================================================================
DBIM A.5.6 CONTENT COMPLIANCE CHECKER
====================================================================

Checks:
23. Content complete and up to date
24. Multi-level CMS workflow
25. Uniform titles usage
26. Accessible PDFs only
27. Date format validation
28. HTTPS & external links validation
29. Archival section validation
30. Ministerial seniority ordering

FEATURES
✔ Automatic Website Scan
✔ Screenshot Capture
✔ DOCX Report Generation
✔ PASS/FAIL Summary
✔ Broken Link Detection
✔ PDF Validation
✔ Evidence Screenshots

USAGE:
    python3 dbim_a5_checker.py
    python3 dbim_a5_checker.py https://example.gov.in

INSTALL:
    pip install selenium webdriver-manager beautifulsoup4
    pip install requests pillow python-docx PyPDF2

====================================================================
"""

import os
import re
import io
import sys
import time
import requests

from urllib.parse import urljoin, urlparse

from bs4 import BeautifulSoup
from PIL import Image

from docx import Document
from docx.shared import Inches

from PyPDF2 import PdfReader

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options

from webdriver_manager.chrome import ChromeDriverManager


# ============================================================
# CONFIG
# ============================================================

OUTPUT_DIR = "dbim_a5_report"

PLACEHOLDER_WORDS = [
    "lorem ipsum",
    "coming soon",
    "tbd",
    "under construction"
]

TITLE_PATTERNS = [
    r"\bDr\.",
    r"\bShri\b",
    r"\bSmt\.",
    r"\bMr\.",
    r"\bMs\."
]

EDITABLE_FORMATS = [
    ".doc",
    ".docx",
    ".ppt",
    ".pptx",
    ".xls",
    ".xlsx"
]

DATE_PATTERNS_VALID = [
    r"\b\d{2}-\d{2}-\d{4}\b",
    r"\b\d{2}/\d{2}/\d{4}\b",
    r"\b\d{1,2}\s+[A-Za-z]+\s+\d{4}\b"
]

DATE_PATTERNS_INVALID = [
    r"\b\d{2}-\d{2}-\d{4}\b"
]


# ============================================================
# OUTPUT FOLDER
# ============================================================

os.makedirs(OUTPUT_DIR, exist_ok=True)


# ============================================================
# DOCX REPORT
# ============================================================

doc = Document()

doc.add_heading(
    "DBIM A.5.6 Content Compliance Report",
    level=0
)

results = []


# ============================================================
# DRIVER SETUP
# ============================================================

def setup_driver():

    options = Options()

    options.add_argument("--headless=new")
    options.add_argument("--disable-gpu")
    options.add_argument("--window-size=1920,5000")

    driver = webdriver.Chrome(
        service=Service(ChromeDriverManager().install()),
        options=options
    )

    return driver


# ============================================================
# UTILITIES
# ============================================================

def get_soup(driver):
    return BeautifulSoup(driver.page_source, "html.parser")


def save_screenshot(driver, name):

    path = os.path.join(
        OUTPUT_DIR,
        f"{name}.png"
    )

    driver.save_screenshot(path)

    return path


def add_result(test_id, title, status, reason, screenshot=None):

    results.append({
        "id": test_id,
        "title": title,
        "status": status,
        "reason": reason
    })

    print(f"\n[{test_id}] {status}")
    print(reason)

    # =====================================================
    # NORMAL DOC FORMAT
    # =====================================================

    doc.add_heading(
        f"{test_id} - {title}",
        level=1
    )

    p = doc.add_paragraph()

    run = p.add_run("Case: ")
    run.bold = True
    p.add_run(f"{test_id} - {title}")

    p = doc.add_paragraph()

    run = p.add_run("Result: ")
    run.bold = True
    p.add_run(status)

    p = doc.add_paragraph()

    run = p.add_run("Reason: ")
    run.bold = True
    p.add_run(reason)

    p = doc.add_paragraph()

    run = p.add_run("Actual: ")
    run.bold = True

    if status == "PASS":
        p.add_run("Requirement satisfied on website.")

    elif status == "FAIL":
        p.add_run("Requirement not satisfied on website.")

    else:
        p.add_run("Manual verification required.")

    p = doc.add_paragraph()

    run = p.add_run("Screenshot Evidence:")
    run.bold = True

    if screenshot and os.path.exists(screenshot):

        doc.add_picture(
            screenshot,
            width=Inches(6)
        )

    doc.add_paragraph(
        "------------------------------------------------------------"
    )


# ============================================================
# TEST 23
# CONTENT COMPLETE & UPDATED
# ============================================================

def test_content_complete(driver, soup):

    screenshot = save_screenshot(
        driver,
        "content_complete"
    )

    text = soup.get_text(" ", strip=True).lower()

    found = []

    for word in PLACEHOLDER_WORDS:

        if word in text:
            found.append(word)

    if found:

        add_result(
            "DBIM-A5-23",
            "Content Complete and Up to Date",
            "FAIL",
            f"Placeholder content found: {', '.join(found)}",
            screenshot
        )

    else:

        add_result(
            "DBIM-A5-23",
            "Content Complete and Up to Date",
            "PASS",
            "No placeholder/incomplete content detected.",
            screenshot
        )


# ============================================================
# TEST 24
# CMS WORKFLOW
# ============================================================

def test_cms_workflow(driver, soup):

    screenshot = save_screenshot(
        driver,
        "cms_workflow"
    )

    text = soup.get_text(" ", strip=True).lower()

    workflow_words = [
        "draft",
        "review",
        "published",
        "approved"
    ]

    found = []

    for w in workflow_words:

        if w in text:
            found.append(w)

    if len(found) >= 2:

        add_result(
            "DBIM-A5-24",
            "CMS Multi-Level Workflow",
            "PASS",
            f"Workflow indicators detected: {', '.join(found)}",
            screenshot
        )

    else:

        add_result(
            "DBIM-A5-24",
            "CMS Multi-Level Workflow",
            "MANUAL CHECK REQUIRED",
            "CMS workflow cannot be fully validated publicly.",
            screenshot
        )


# ============================================================
# TEST 25
# TITLES VALIDATION
# ============================================================

def test_titles(driver, soup):

    screenshot = save_screenshot(
        driver,
        "titles_validation"
    )

    text = soup.get_text(" ", strip=True)

    found_titles = []

    for pattern in TITLE_PATTERNS:

        matches = re.findall(pattern, text)

        if matches:
            found_titles.extend(matches)

    if found_titles:

        add_result(
            "DBIM-A5-25",
            "Uniform Titles Usage",
            "PASS",
            f"Titles detected: {', '.join(set(found_titles))}",
            screenshot
        )

    else:

        add_result(
            "DBIM-A5-25",
            "Uniform Titles Usage",
            "FAIL",
            "No official titles detected.",
            screenshot
        )


# ============================================================
# TEST 26
# PDF VALIDATION
# ============================================================

def test_pdfs(driver, soup, base_url):

    screenshot = save_screenshot(
        driver,
        "pdf_validation"
    )

    anchors = soup.find_all("a", href=True)

    editable_found = []
    pdf_found = 0

    for a in anchors:

        href = a["href"].lower()

        for ext in EDITABLE_FORMATS:

            if href.endswith(ext):
                editable_found.append(href)

        if href.endswith(".pdf"):
            pdf_found += 1

    if editable_found:

        add_result(
            "DBIM-A5-26",
            "Accessible PDFs Only",
            "FAIL",
            f"Editable documents found: {len(editable_found)}",
            screenshot
        )

    elif pdf_found > 0:

        add_result(
            "DBIM-A5-26",
            "Accessible PDFs Only",
            "PASS",
            f"{pdf_found} PDF files detected and no editable files found.",
            screenshot
        )

    else:

        add_result(
            "DBIM-A5-26",
            "Accessible PDFs Only",
            "MANUAL CHECK REQUIRED",
            "No downloadable documents detected.",
            screenshot
        )


# ============================================================
# TEST 27
# DATE FORMAT
# ============================================================

def test_date_format(driver, soup):

    screenshot = save_screenshot(
        driver,
        "date_format"
    )

    text = soup.get_text(" ", strip=True)

    valid_found = False

    for pattern in DATE_PATTERNS_VALID:

        if re.search(pattern, text):
            valid_found = True
            break

    if valid_found:

        add_result(
            "DBIM-A5-27",
            "Date Format Validation",
            "PASS",
            "Day-before-month date format detected.",
            screenshot
        )

    else:

        add_result(
            "DBIM-A5-27",
            "Date Format Validation",
            "FAIL",
            "Required date format not detected.",
            screenshot
        )


# ============================================================
# TEST 28
# HTTPS LINKS
# ============================================================

def test_https_links(driver, soup):

    screenshot = save_screenshot(
        driver,
        "https_links"
    )

    anchors = soup.find_all("a", href=True)

    insecure = []
    broken = []

    for a in anchors:

        href = a["href"]

        if href.startswith("http://"):
            insecure.append(href)

        if href.startswith("http"):

            try:

                r = requests.get(
                    href,
                    timeout=10
                )

                if r.status_code >= 400:
                    broken.append(href)

            except:
                broken.append(href)

    if insecure or broken:

        reason = ""

        if insecure:
            reason += f"Insecure links: {len(insecure)}. "

        if broken:
            reason += f"Broken links: {len(broken)}."

        add_result(
            "DBIM-A5-28",
            "HTTPS & Link Validation",
            "FAIL",
            reason,
            screenshot
        )

    else:

        add_result(
            "DBIM-A5-28",
            "HTTPS & Link Validation",
            "PASS",
            "All checked links are HTTPS and functional.",
            screenshot
        )


# ============================================================
# TEST 29
# ARCHIVE SECTION
# ============================================================

def test_archive(driver, soup):

    screenshot = save_screenshot(
        driver,
        "archive_section"
    )

    text = soup.get_text(" ", strip=True).lower()

    keywords = [
        "archive",
        "archival",
        "older notices",
        "past notices"
    ]

    found = False

    for k in keywords:

        if k in text:
            found = True
            break

    if found:

        add_result(
            "DBIM-A5-29",
            "Archive Section",
            "PASS",
            "Archive section detected.",
            screenshot
        )

    else:

        add_result(
            "DBIM-A5-29",
            "Archive Section",
            "FAIL",
            "Archive section not detected.",
            screenshot
        )


# ============================================================
# TEST 30
# SENIORITY ORDER
# ============================================================

def test_seniority(driver, soup):

    screenshot = save_screenshot(
        driver,
        "seniority"
    )

    text = soup.get_text(" ", strip=True).lower()

    minister = text.find("minister")
    mos = text.find("minister of state")

    if minister != -1 and mos != -1:

        if minister < mos:

            add_result(
                "DBIM-A5-30",
                "Ministerial Seniority Order",
                "PASS",
                "Minister appears before Minister of State.",
                screenshot
            )

        else:

            add_result(
                "DBIM-A5-30",
                "Ministerial Seniority Order",
                "FAIL",
                "Ministerial hierarchy order incorrect.",
                screenshot
            )

    else:

        add_result(
            "DBIM-A5-30",
            "Ministerial Seniority Order",
            "MANUAL CHECK REQUIRED",
            "Unable to automatically determine hierarchy.",
            screenshot
        )


# ============================================================
# SUMMARY
# ============================================================

def add_summary():

    doc.add_page_break()

    doc.add_heading(
        "Summary",
        level=1
    )

    table = doc.add_table(
        rows=1,
        cols=4
    )

    hdr = table.rows[0].cells

    hdr[0].text = "Test ID"
    hdr[1].text = "Test"
    hdr[2].text = "Status"
    hdr[3].text = "Reason"

    for r in results:

        row = table.add_row().cells

        row[0].text = r["id"]
        row[1].text = r["title"]
        row[2].text = r["status"]
        row[3].text = r["reason"]


# ============================================================
# MAIN
# ============================================================

def main():

    if len(sys.argv) > 1:
        url = sys.argv[1]

    else:
        url = input(
            "Enter Website URL: "
        ).strip()

    if not url.startswith("http"):
        url = "https://" + url

    print("\n================================================")
    print("DBIM A.5.6 CONTENT COMPLIANCE CHECKER")
    print("================================================")
    print("Target:", url)

    driver = setup_driver()

    try:

        driver.get(url)

        time.sleep(5)

        homepage_ss = save_screenshot(
            driver,
            "homepage"
        )

        doc.add_heading(
            "Homepage Screenshot",
            level=1
        )

        doc.add_picture(
            homepage_ss,
            width=Inches(6)
        )

        soup = get_soup(driver)

        # TESTS

        test_content_complete(driver, soup)

        test_cms_workflow(driver, soup)

        test_titles(driver, soup)

        test_pdfs(driver, soup, url)

        test_date_format(driver, soup)

        test_https_links(driver, soup)

        test_archive(driver, soup)

        test_seniority(driver, soup)

        # SUMMARY

        add_summary()

        report_path = os.path.join(
            OUTPUT_DIR,
            "DBIM_A5_Compliance_Report.docx"
        )

        doc.save(report_path)

        print("\n================================================")
        print("REPORT GENERATED")
        print("================================================")
        print("DOCX REPORT:", report_path)
        print("SCREENSHOTS:", OUTPUT_DIR)

    except Exception as e:

        print("\nERROR:", str(e))

    finally:

        driver.quit()


if __name__ == "__main__":
    main()