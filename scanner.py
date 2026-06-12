

import sys
import os
import re
import json
import time
import argparse
from datetime import datetime
from urllib.parse import urlparse

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import shutil

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
NOTO_SANS_VARIANTS = [
    "noto sans", "notosans", "noto-sans",
    "noto sans devanagari", "noto sans tamil",
    "noto sans bengali", "noto sans telugu",
    "noto sans kannada", "noto sans malayalam",
    "noto sans gujarati", "noto sans gurmukhi",
    "noto sans oriya", "noto sans sinhala",
]

ELEMENTS_TO_CHECK = {
    "body":   "Body / Global",
    "h1":     "Heading 1",
    "h2":     "Heading 2",
    "h3":     "Heading 3",
    "p":      "Paragraph",
    "a":      "Anchor / Link",
    "li":     "List Item",
    "span":   "Span",
    "button": "Button",
    "input":  "Input Field",
    "nav":    "Navigation",
    "footer": "Footer",
    "header": "Header",
    "table":  "Table",
    "label":  "Label",
}

# Add a small dictionary of common Hinglish terms for basic filtering
HINGLISH_DICTIONARY = [
    "hai", "haan", "nahi", "kya", "kyun", "kaise", "kab", "kahan",
    "bhi", "toh", "sirf", "lekin", "aur", "ya", "par", "mein"
]


# ---------------------------------------------------------------------------
# Browser helpers
# ---------------------------------------------------------------------------
def _find_chrome_binary():
    """Locate a Chrome/Chromium binary on the system."""
    for name in ["google-chrome", "google-chrome-stable", "chromium-browser", "chromium"]:
        path = shutil.which(name)
        if path:
            return path
    return None


def create_driver():
    """Create an anti-bot-detection headless Chrome driver."""
    opts = Options()
    chrome_bin = _find_chrome_binary()
    if chrome_bin:
        opts.binary_location = chrome_bin

    opts.add_argument("--headless=new")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--window-size=1920,1080")
    opts.add_argument("--disable-gpu")
    opts.add_argument("--ignore-certificate-errors")
    opts.add_argument("--disable-extensions")
    opts.add_argument("--disable-software-rasterizer")

    # ---- Anti-bot-detection ----
    opts.add_argument("--disable-blink-features=AutomationControlled")
    opts.add_argument(
        "--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
    opts.add_experimental_option("excludeSwitches", ["enable-automation"])
    opts.add_experimental_option("useAutomationExtension", False)

    driver = webdriver.Chrome(options=opts)
    driver.set_page_load_timeout(60)

    # Hide navigator.webdriver flag so WAF JS checks don't see Selenium
    driver.execute_cdp_cmd(
        "Page.addScriptToEvaluateOnNewDocument",
        {"source": """
            Object.defineProperty(navigator, 'webdriver', {get: () => undefined});
            Object.defineProperty(navigator, 'plugins', {get: () => [1,2,3]});
            Object.defineProperty(navigator, 'languages', {get: () => ['en-IN','en','hi']});
        """}
    )

    return driver


def wait_for_page(driver, timeout=30):
    """Wait until the page has fully loaded."""
    WebDriverWait(driver, timeout).until(
        lambda d: d.execute_script("return document.readyState") == "complete"
    )
    time.sleep(2)  # extra settle time for web-fonts


# ---------------------------------------------------------------------------
# Font, Layout, Content & Contrast Detection Methods
# ---------------------------------------------------------------------------
def is_noto_sans(font_string: str) -> bool:
    if not font_string:
        return False
    lower = font_string.lower()
    return any(v in lower for v in NOTO_SANS_VARIANTS)


def get_computed_font(driver, css_selector: str):
    script = """
    var el = document.querySelector(arguments[0]);
    if (!el) return null;
    var cs = window.getComputedStyle(el);
    return cs.getPropertyValue('font-family');
    """
    return driver.execute_script(script, css_selector)


def check_font_import_in_css(driver):
    script = """
    var results = {link_tags: [], font_face: [], import_rules: [], native_check: false};

    document.querySelectorAll('link[rel="stylesheet"]').forEach(function(l){
        if (l.href && l.href.toLowerCase().indexOf('noto') !== -1)
            results.link_tags.push(l.href);
    });

    try {
        results.native_check = document.fonts.check("1em 'Noto Sans'");
    } catch(e) {}

    try {
        for (var i = 0; i < document.styleSheets.length; i++) {
            try {
                var rules = document.styleSheets[i].cssRules || [];
                for (var j = 0; j < rules.length; j++) {
                    var r = rules[j];
                    if (r.type === CSSRule.FONT_FACE_RULE) {
                        var ff = r.style.getPropertyValue('font-family') || '';
                        if (ff.toLowerCase().indexOf('noto') !== -1)
                            results.font_face.push(ff);
                    }
                    if (r.type === CSSRule.IMPORT_RULE && r.href) {
                        if (r.href.toLowerCase().indexOf('noto') !== -1)
                            results.import_rules.push(r.href);
                    }
                }
            } catch(e) { }
        }
    } catch(e) {}
    return JSON.stringify(results);
    """
    raw = driver.execute_script(script)
    return json.loads(raw) if raw else {}


def check_inline_styles_for_noto(driver):
    script = """
    var found = [];
    document.querySelectorAll('[style]').forEach(function(el){
        var s = el.getAttribute('style') || '';
        if (s.toLowerCase().indexOf('noto') !== -1)
            found.push({tag: el.tagName, style: s.substring(0, 200)});
    });
    return JSON.stringify(found);
    """
    raw = driver.execute_script(script)
    return json.loads(raw) if raw else []


def check_body_alignment(driver):
    script = """
    var body = document.querySelector('body');
    if (!body) return JSON.stringify({ passed: true, reason: 'No body tag found.' });
    var style = window.getComputedStyle(body);
    var align = style.textAlign;
    if (align === 'left' || align === 'start' || align === '-webkit-auto' ||
        align === '-webkit-left' || align === 'match-parent') {
        return JSON.stringify({ passed: true,
            reason: 'Body text is left-aligned (found: ' + align + ').' });
    } else {
        return JSON.stringify({ passed: false,
            reason: 'Body text is not left-aligned (found: ' + align + ').' });
    }
    """
    raw = driver.execute_script(script)
    return json.loads(raw) if raw else {"passed": False, "reason": "Failed to execute script"}


def check_table_alignment(driver):
    script = """
    var tables = document.querySelectorAll('table');
    if (tables.length === 0)
        return JSON.stringify({ passed: true, reason: 'No tables found on the page.' });

    var issues = [];

    for (var i = 0; i < tables.length; i++) {
        var table = tables[i];

        var ths = table.querySelectorAll('th');
        for (var j = 0; j < ths.length; j++) {
            var th = ths[j];
            var style = window.getComputedStyle(th);
            var align = style.textAlign;
            if (align !== 'center' && align !== '-webkit-center') {
                issues.push('<th> element not center-aligned (found: ' + align + ')');
            }
        }

        var tds = table.querySelectorAll('td');
        for (var k = 0; k < tds.length; k++) {
            var td = tds[k];
            var text = td.innerText.trim();
            if (!text) continue;

            var style = window.getComputedStyle(td);
            var align = style.textAlign;

            var cleanText = text.replace(/[\\u20B9\\$\\%,\\s]/g, '');
            var isNumber = cleanText.length > 0 &&
                           !isNaN(Number(cleanText)) && isFinite(cleanText);

            if (isNumber) {
                if (align !== 'right' && align !== '-webkit-right' && align !== 'end') {
                    issues.push('<td> with numeric data not right-aligned (found: ' +
                        align + ', text: "' + text.substring(0, 10) + '")');
                }
            } else {
                if (align !== 'left' && align !== 'start' && align !== '-webkit-left' &&
                    align !== '-webkit-match-parent' && align !== 'match-parent') {
                    issues.push('<td> with text data not left-aligned (found: ' +
                        align + ', text: "' + text.substring(0, 10) + '")');
                }
            }
        }
    }

    if (issues.length > 0) {
        var reason = issues.slice(0, 3).join('; ') +
            (issues.length > 3 ? ' (and ' + (issues.length - 3) + ' more...)' : '');
        return JSON.stringify({ passed: false, reason: reason });
    }

    return JSON.stringify({ passed: true,
        reason: 'All ' + tables.length + ' table(s) alignments comply with guidelines.' });
    """
    raw = driver.execute_script(script)
    return json.loads(raw) if raw else {"passed": False, "reason": "Failed to execute script"}


# ---------------------------------------------------------------------------
# TC: Type Scale - redesigned to use querySelectorAll and full-page analysis
# ---------------------------------------------------------------------------
def check_type_scale(driver):
    script = """
    function getAllFontSizes(selector) {
        var els = document.querySelectorAll(selector);
        var sizes = [];
        for (var i = 0; i < els.length; i++) {
            var el = els[i];
            if (el.offsetParent === null && el.tagName !== 'BODY') continue;
            var sizeStr = window.getComputedStyle(el).getPropertyValue('font-size');
            var px = parseFloat(sizeStr);
            if (!isNaN(px) && px > 0) {
                sizes.push(Math.round(px * 100) / 100);
            }
        }
        return sizes;
    }

    function uniqueSorted(arr) {
        var seen = {};
        var out = [];
        for (var i = 0; i < arr.length; i++) {
            if (!seen[arr[i]]) { seen[arr[i]] = true; out.push(arr[i]); }
        }
        out.sort(function(a, b) { return b - a; });
        return out;
    }

    function median(arr) {
        if (arr.length === 0) return null;
        var sorted = arr.slice().sort(function(a, b) { return a - b; });
        var mid = Math.floor(sorted.length / 2);
        return sorted.length % 2 !== 0
            ? sorted[mid]
            : Math.round(((sorted[mid - 1] + sorted[mid]) / 2) * 100) / 100;
    }

    function maxVal(arr) {
        if (arr.length === 0) return null;
        return Math.max.apply(null, arr);
    }

    var h1Sizes = getAllFontSizes('h1');
    var h2Sizes = getAllFontSizes('h2');
    var h3Sizes = getAllFontSizes('h3');
    var pSizes  = getAllFontSizes('p');

    var maxH1 = maxVal(h1Sizes);
    var maxH2 = maxVal(h2Sizes);
    var maxH3 = maxVal(h3Sizes);
    var medP  = median(pSizes);
    var maxP  = maxVal(pSizes);

    var uniqH1 = uniqueSorted(h1Sizes);
    var uniqH2 = uniqueSorted(h2Sizes);
    var uniqH3 = uniqueSorted(h3Sizes);

    var issues = [];
    var info   = [];

    // 1. Hierarchy checks
    if (maxH1 !== null && maxH2 !== null) {
        if (maxH1 === maxH2) {
            issues.push('H1 (' + maxH1 + 'px) is equal to H2 (' + maxH2 +
                        'px), violating H1 > H2 hierarchy.');
        } else if (maxH1 < maxH2) {
            issues.push('H1 (' + maxH1 + 'px) is smaller than H2 (' + maxH2 +
                        'px), violating H1 > H2 hierarchy.');
        }
    }

    if (maxH2 !== null && maxH3 !== null) {
        if (maxH2 === maxH3) {
            issues.push('H2 (' + maxH2 + 'px) is equal to H3 (' + maxH3 +
                        'px), violating H2 > H3 hierarchy.');
        } else if (maxH2 < maxH3) {
            issues.push('H2 (' + maxH2 + 'px) is smaller than H3 (' + maxH3 +
                        'px), violating H2 > H3 hierarchy.');
        }
    }

    if (maxH3 !== null && medP !== null) {
        if (maxH3 <= medP) {
            issues.push('H3 (' + maxH3 + 'px) is not larger than the median paragraph size (' +
                        medP + 'px), violating H3 > P hierarchy.');
        }
    }

    // 2. Paragraph larger than headings
    if (maxP !== null && maxH3 !== null && maxP > maxH3) {
        issues.push('Paragraph elements found up to ' + maxP + 'px, exceeding H3 (' +
                    maxH3 + 'px) - paragraph must not be larger than body headings.');
    }

    if (maxP !== null && maxH2 !== null && maxP >= maxH2) {
        issues.push('Paragraph elements found up to ' + maxP +
                    'px, equalling or exceeding H2 (' + maxH2 +
                    'px) - this is a critical hierarchy violation.');
    }

    // 3. Inconsistent heading sizes across the page
    if (uniqH1.length > 1) {
        info.push('Inconsistent H1 sizes detected across the page: ' +
                  uniqH1.join('px, ') + 'px.');
    }
    if (uniqH2.length > 1) {
        issues.push('Inconsistent H2 sizes detected across the page: ' +
                    uniqH2.join('px, ') + 'px.');
    }
    if (uniqH3.length > 1) {
        issues.push('Inconsistent H3 sizes detected across the page: ' +
                    uniqH3.join('px, ') + 'px.');
    }

    // 4. Missing elements - informational only
    if (h1Sizes.length === 0) info.push('No visible H1 elements found - H1 check skipped.');
    if (h2Sizes.length === 0) info.push('No visible H2 elements found - H2 check skipped.');
    if (h3Sizes.length === 0) info.push('No visible H3 elements found - H3 check skipped.');
    if (pSizes.length  === 0) info.push('No visible P elements found - paragraph check skipped.');

    // 5. Build result
    if (issues.length > 0) {
        var summary = 'Typography hierarchy is inconsistent across the page. ';
        summary += issues.join(' | ');
        if (info.length > 0) summary += ' [INFO: ' + info.join(' ') + ']';
        return JSON.stringify({ passed: false, reason: summary });
    }

    var passParts = [];
    if (maxH1 !== null) passParts.push('H1=' + maxH1 + 'px (' + h1Sizes.length + ' el)');
    if (maxH2 !== null) passParts.push('H2=' + maxH2 + 'px (' + h2Sizes.length + ' el)');
    if (maxH3 !== null) passParts.push('H3=' + maxH3 + 'px (' + h3Sizes.length + ' el)');
    if (medP  !== null) passParts.push('P(median)=' + medP + 'px (' + pSizes.length + ' el)');

    var passReason = 'Heading scale hierarchy (H1 > H2 > H3 > P) maintained. ' +
                     'Evaluated all visible elements - ' + passParts.join(', ') + '.';
    if (info.length > 0) passReason += ' [INFO: ' + info.join(' ') + ']';

    return JSON.stringify({ passed: true, reason: passReason });
    """
    raw = driver.execute_script(script)
    return json.loads(raw) if raw else {"passed": False, "reason": "Failed to execute script"}


def check_capital_case_and_hinglish(driver):
    try:
        paragraphs = driver.find_elements(By.TAG_NAME, "p")
    except Exception:
        return {"passed": True,
                "reason": "No <p> tags found to analyze for capitalization or Hinglish."}

    issues = []

    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue

        words = text.split()
        if len(words) > 5:
            # 1. Capital Case Check
            alpha_only = re.sub(r'[^a-zA-Z]', '', text)
            if len(alpha_only) > 0 and alpha_only.isupper():
                issues.append(
                    f"Found completely capitalized long paragraph: '{text[:40]}...'"
                )

        # 2. Hinglish Content Check
        lower_words = [w.lower() for w in words]
        for w in lower_words:
            clean_word = re.sub(r'[^\w\s]', '', w)
            if clean_word in HINGLISH_DICTIONARY:
                issues.append(
                    f"Potential Hinglish detected ('{clean_word}') in: '{text[:40]}...'"
                )
                break

    if len(issues) > 0:
        reason = issues[0] + (
            f" (and {len(issues)-1} other issues)" if len(issues) > 1 else ""
        )
        return {"passed": False, "reason": reason}

    return {"passed": True,
            "reason": "No capitalized long sentences or explicit Hinglish content detected."}


def check_color_contrast(driver):
    script = """
    function parseRGB(rgbString) {
        var match = rgbString.match(/rgba?\\((\\d+),\\s*(\\d+),\\s*(\\d+)/);
        if (!match) return null;
        return [parseInt(match[1]), parseInt(match[2]), parseInt(match[3])];
    }

    function getLuminance(rgb) {
        var a = rgb.map(function (v) {
            v /= 255;
            return v <= 0.03928 ? v / 12.92 : Math.pow((v + 0.055) / 1.055, 2.4);
        });
        return a[0] * 0.2126 + a[1] * 0.7152 + a[2] * 0.0722;
    }

    function getContrastRatio(rgb1, rgb2) {
        var lum1 = getLuminance(rgb1);
        var lum2 = getLuminance(rgb2);
        var brightest = Math.max(lum1, lum2);
        var darkest = Math.min(lum1, lum2);
        return (brightest + 0.05) / (darkest + 0.05);
    }

    var elements = document.querySelectorAll('p, h1, h2, h3, h4, h5, h6, a, button, span');
    var issues = [];
    var checkedCount = 0;

    for (var i = 0; i < elements.length; i++) {
        var el = elements[i];
        if (el.innerText.trim() === '') continue;

        var style = window.getComputedStyle(el);
        var color = style.getPropertyValue('color');
        var bg    = style.getPropertyValue('background-color');

        var parent = el;
        while (bg === 'rgba(0, 0, 0, 0)' || bg === 'transparent') {
            parent = parent.parentElement;
            if (!parent) { bg = 'rgb(255, 255, 255)'; break; }
            bg = window.getComputedStyle(parent).getPropertyValue('background-color');
        }

        var colorRGB = parseRGB(color);
        var bgRGB    = parseRGB(bg);

        if (colorRGB && bgRGB) {
            var ratio = getContrastRatio(colorRGB, bgRGB);
            if (ratio < 4.5) {
                var sample = el.innerText.substring(0, 25).replace(/[\\n\\r]/g, ' ');
                issues.push("Low contrast (" + ratio.toFixed(2) +
                            ":1) on text: '" + sample + "...'");
                if (issues.length >= 3) break;
            }
            checkedCount++;
        }
    }

    if (issues.length > 0) {
        return JSON.stringify({ passed: false,
            reason: "Contrast Failures: " + issues.join(" | ") });
    }
    return JSON.stringify({ passed: true,
        reason: "Checked " + checkedCount + " elements. All met 4.5:1 standard." });
    """
    raw = driver.execute_script(script)
    return json.loads(raw) if raw else {"passed": False, "reason": "Failed to execute script."}


# ---------------------------------------------------------------------------
# TC-16: Button padding consistency (DBIM 4.5)
# ---------------------------------------------------------------------------
def check_button_padding_consistency(driver):
    script = """
    var buttons = document.querySelectorAll(
        'button, input[type="submit"], input[type="button"], a[role="button"]'
    );
    if (buttons.length === 0) {
        return JSON.stringify({passed: true, reason: 'No buttons found on page.'});
    }

    var paddings = [];
    for (var i = 0; i < buttons.length; i++) {
        var s  = window.getComputedStyle(buttons[i]);
        var pt = Math.round(parseFloat(s.paddingTop));
        var pb = Math.round(parseFloat(s.paddingBottom));
        var pl = Math.round(parseFloat(s.paddingLeft));
        var pr = Math.round(parseFloat(s.paddingRight));
        paddings.push(pt + ',' + pb + ',' + pl + ',' + pr);
    }

    var unique = paddings.filter(function(v, i, a){ return a.indexOf(v) === i; });
    if (unique.length > 3) {
        return JSON.stringify({
            passed: false,
            reason: 'Inconsistent button padding found (' + unique.length +
                    ' distinct padding values across ' + buttons.length + ' buttons). ' +
                    'Values: ' + unique.slice(0, 4).join(' | ')
        });
    }
    return JSON.stringify({
        passed: true,
        reason: 'Button padding is consistent across ' + buttons.length +
                ' button(s). Distinct patterns: ' + unique.join(' | ')
    });
    """
    raw = driver.execute_script(script)
    return json.loads(raw) if raw else {"passed": False, "reason": "Failed to execute script."}


# ---------------------------------------------------------------------------
# TC-17: Button state differentiation - hover / focus / disabled (DBIM 4.5)
# ---------------------------------------------------------------------------
def check_button_states(driver):
    script = """
    var buttons = document.querySelectorAll(
        'button:not([disabled]), input[type="submit"]:not([disabled]), a[role="button"]'
    );
    if (buttons.length === 0) {
        return JSON.stringify({passed: true, reason: 'No enabled buttons found.'});
    }

    var issues = [], seen = {};
    var limit = Math.min(buttons.length, 10);

    for (var i = 0; i < limit; i++) {
        var btn = buttons[i];
        var label = (btn.innerText || btn.value ||
                     btn.getAttribute('aria-label') || btn.tagName)
                    .trim().substring(0, 30);
        if (seen[label]) continue;
        seen[label] = true;

        var base       = window.getComputedStyle(btn);
        var baseBg     = base.backgroundColor;
        var baseColor  = base.color;
        var baseBorder = base.borderColor;
        var baseOutline= base.outline;

        // -- Hover --
        try { btn.dispatchEvent(new MouseEvent('mouseover', {bubbles:true})); } catch(e){}
        var hov = window.getComputedStyle(btn);
        var hoverChanged = (hov.backgroundColor !== baseBg ||
                            hov.color !== baseColor ||
                            hov.borderColor !== baseBorder);
        if (!hoverChanged) {
            issues.push('No hover style change on "' + label + '"');
        }

        // -- Focus --
        try { btn.focus(); } catch(e){}
        var foc = window.getComputedStyle(btn);
        var focusChanged = (foc.outline !== baseOutline ||
                            foc.backgroundColor !== baseBg ||
                            foc.borderColor !== baseBorder);
        if (!focusChanged) {
            issues.push('No focus style change on "' + label + '"');
        }
        try { btn.blur(); } catch(e){}

        if (issues.length >= 4) break;
    }

    // -- Disabled state --
    var disabled = document.querySelectorAll(
        'button[disabled], input[type="submit"][disabled], input[type="button"][disabled]'
    );
    if (disabled.length === 0) {
        issues.push('No disabled button found to verify disabled state styling.');
    }

    return JSON.stringify({
        passed: issues.length === 0,
        reason: issues.length
            ? issues.join(' | ')
            : 'Hover, focus, and disabled button states all have distinct styles.'
    });
    """
    raw = driver.execute_script(script)
    return json.loads(raw) if raw else {"passed": False, "reason": "Failed to execute script."}


# ---------------------------------------------------------------------------
# TC-18: Mouse hover - noticeable change on clickable items (DBIM 4.5)
# ---------------------------------------------------------------------------
def check_hover_feedback(driver):
    script = """
    var els = document.querySelectorAll(
        'a, button, input[type="button"], input[type="submit"]'
    );
    if (els.length === 0) {
        return JSON.stringify({passed: true, reason: 'No clickable elements found.'});
    }
    var issues = [], seen = {};
    for (var i = 0; i < els.length; i++) {
        var el  = els[i];
        var s0  = window.getComputedStyle(el);
        var b0  = s0.backgroundColor;
        var c0  = s0.color;
        var br0 = s0.borderColor;
        try { el.dispatchEvent(new MouseEvent('mouseover', {bubbles:true})); } catch(e){}
        var s1 = window.getComputedStyle(el);
        var changed = b0 !== s1.backgroundColor ||
                      c0 !== s1.color ||
                      br0 !== s1.borderColor;
        if (!changed) {
            var label = (el.innerText || el.value || el.getAttribute('aria-label') ||
                         el.getAttribute('title') || el.href || el.tagName)
                        .trim().substring(0, 40);
            if (label === '\u2039' || label === '\u203a') continue;
            var msg = 'No hover feedback: "' + label + '"';
            if (!seen[msg]) { issues.push(msg); seen[msg] = true; }
        }
        if (issues.length >= 5) break;
    }
    return JSON.stringify({
        passed: issues.length === 0,
        reason: issues.length
            ? issues.join(' | ')
            : 'All clickable elements show a noticeable change on mouse hover.'
    });
    """
    raw = driver.execute_script(script)
    return json.loads(raw) if raw else {"passed": False, "reason": "Failed to execute script."}


# ---------------------------------------------------------------------------
# Screenshot helpers
# ---------------------------------------------------------------------------
def take_full_screenshot(driver, path, max_width=1600, max_height=8000):
    """
    Capture a full-page screenshot and resize it so python-docx can embed it safely.
    Falls back to a normal viewport screenshot if anything goes wrong.
    """
    try:
        # Try to get full page height
        total_height = driver.execute_script(
            "return Math.max(document.body.scrollHeight, document.documentElement.scrollHeight);"
        )
        total_width = driver.execute_script(
            "return Math.max(document.body.scrollWidth, document.documentElement.scrollWidth);"
        )
        original_size = driver.get_window_size()

        # Cap dimensions so the PNG doesn't explode in size
        cap_h = min(total_height, max_height)
        cap_w = min(total_width, max_width)
        driver.set_window_size(cap_w, cap_h)
        time.sleep(0.5)

        driver.save_screenshot(path)

        # Restore original size
        driver.set_window_size(original_size["width"], original_size["height"])
    except Exception:
        # Fallback: plain viewport screenshot
        driver.save_screenshot(path)

    # Post-process: downscale if image is still too big for python-docx
    try:
        from PIL import Image
        img = Image.open(path)
        w, h = img.size
        if w > max_width or h > max_height:
            ratio = min(max_width / w, max_height / h)
            new_size = (int(w * ratio), int(h * ratio))
            img = img.resize(new_size, Image.LANCZOS)
            img.save(path, optimize=True)
        img.close()
    except Exception:
        pass  # if Pillow not available, leave the file as-is

    return path


def take_element_screenshot(driver, css_selector, path):
    try:
        el = driver.find_element(By.CSS_SELECTOR, css_selector)
        el.screenshot(path)
    except Exception:
        driver.save_screenshot(path)
    return path


# ---------------------------------------------------------------------------
# Audit engine
# ---------------------------------------------------------------------------
class AuditResult:
    def __init__(self, tc_id, name, description, result, reason, screenshot=None):
        self.tc_id       = tc_id
        self.name        = name
        self.description = description
        self.result      = result
        self.reason      = reason
        self.screenshot  = screenshot


def run_audit(url: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)
    ss_dir = os.path.join(output_dir, "screenshots")
    os.makedirs(ss_dir, exist_ok=True)

    print(f"[*] Launching browser for {url} ...")
    driver = create_driver()
    results = []
    tc_counter = 0

    try:
        driver.get(url)
        wait_for_page(driver)

        # -- TC-00: Full page screenshot (Base Desktop View) --------------------
        full_ss = take_full_screenshot(
            driver, os.path.join(ss_dir, "full_page.png")
        )

        # -- Guard: detect WAF / Access Denied pages ----------------------------
        page_text = driver.page_source.lower()
        block_signals = [
            "access denied", "403 forbidden", "you don't have permission",
            "akamai", "cloudflare", "bot detection", "captcha",
            "ddos-guard", "enable cookies",
        ]
        if any(sig in page_text for sig in block_signals):
            print("[!] WAF / bot-protection page detected - aborting audit.")
            results.append(AuditResult(
                "BLOCK-01",
                "Bot Protection Detected",
                "Pre-audit check: confirm site is accessible to automated scanner.",
                "FAIL",
                "The site returned an Access Denied / bot-protection page. "
                "All subsequent checks would be meaningless. "
                "Try auditing in visible (non-headless) mode or use a whitelisted IP.",
                full_ss,
            ))
            return results

        # -- TC-01: CSS / link-level Noto Sans loading --------------------------
        tc_counter += 1
        css_info  = check_font_import_in_css(driver)
        has_link  = bool(css_info.get("link_tags"))
        has_face  = bool(css_info.get("font_face"))
        has_import= bool(css_info.get("import_rules"))
        native_ok = css_info.get("native_check", False)
        loaded    = has_link or has_face or has_import or native_ok

        if loaded:
            parts = []
            if has_link:   parts.append(f"<link> tags: {css_info['link_tags']}")
            if has_face:   parts.append(f"@font-face: {css_info['font_face']}")
            if has_import: parts.append(f"@import: {css_info['import_rules']}")
            if native_ok:  parts.append("Native browser layout stack active.")
            reason = "Noto Sans font resources detected. " + "; ".join(parts)
        else:
            reason = "No <link>, @font-face, or @import rule referencing Noto Sans found."

        results.append(AuditResult(
            f"TC-{tc_counter:02d}",
            "Font Resource Loading",
            "Verify that Noto Sans font files are loaded via CSS.",
            "PASS" if loaded else "FAIL",
            reason,
            full_ss,
        ))

        # -- TC-02: Inline style references -------------------------------------
        tc_counter += 1
        inline = check_inline_styles_for_noto(driver)
        if inline:
            reason = f"Found {len(inline)} element(s) with inline styles referencing Noto Sans."
            res = "PASS"
        else:
            reason = "No inline style attribute references Noto Sans."
            res = "INFO"

        results.append(AuditResult(
            f"TC-{tc_counter:02d}",
            "Inline Style Noto Sans Reference",
            "Check whether any element uses Noto Sans via inline style.",
            res, reason, full_ss,
        ))

        # -- TC-03 onwards: Per-element computed font checks --------------------
        for selector, label in ELEMENTS_TO_CHECK.items():
            tc_counter += 1
            font_family = get_computed_font(driver, selector)
            ss_path = os.path.join(ss_dir, f"element_{selector}.png")

            if font_family is None:
                results.append(AuditResult(
                    f"TC-{tc_counter:02d}",
                    f"Computed Font - {label} (<{selector}>)",
                    f"Verify the <{selector}> element uses Noto Sans.",
                    "SKIP",
                    f"No <{selector}> element found on the page.",
                ))
                continue

            take_element_screenshot(driver, selector, ss_path)
            passes = is_noto_sans(font_family)
            reason = f"Computed font-family: \"{font_family}\". "
            reason += ("Noto Sans detected - COMPLIANT."
                       if passes else "Noto Sans NOT detected - NON-COMPLIANT.")

            results.append(AuditResult(
                f"TC-{tc_counter:02d}",
                f"Computed Font - {label} (<{selector}>)",
                f"Verify the <{selector}> element uses Noto Sans.",
                "PASS" if passes else "FAIL",
                reason,
                ss_path,
            ))

        # -- Explicit Layout, Content & Type Scale Rules ------------------------

        # 1. Body Text Layout Check
        body_align = check_body_alignment(driver)
        results.append(AuditResult(
            "ALIGN-01",
            "Body Text Alignment Check",
            "Verify that standard page body text layers render strictly left-aligned.",
            "PASS" if body_align.get("passed") else "FAIL",
            body_align.get("reason", "Unknown extraction error"),
            full_ss,
        ))

        # 2. Complex Table Grid Alignment Check
        table_align = check_table_alignment(driver)
        results.append(AuditResult(
            "ALIGN-02",
            "Table Data Structure Alignment",
            "Verify tables enforce left-aligned text, right-aligned numbers, center headers.",
            "PASS" if table_align.get("passed") else "FAIL",
            table_align.get("reason", "Unknown extraction error"),
            full_ss,
        ))

        # 3. Content Formatting (Capital Case & Hinglish)
        content_check = check_capital_case_and_hinglish(driver)
        results.append(AuditResult(
            "CONTENT-01",
            "Content Casing & Language Purity Check",
            "Verify paragraphs do not use full capital case or contain basic Hinglish.",
            "PASS" if content_check.get("passed") else "FAIL",
            content_check.get("reason", "Unknown extraction error"),
            full_ss,
        ))

        # 4. Type Scale Check (Desktop)
        scale_check_desktop = check_type_scale(driver)
        results.append(AuditResult(
            "SCALE-01",
            "Typographic Hierarchy Scale (Desktop)",
            "Verify font sizes maintain a logical hierarchy (H1>H2>H3>P) at 1920x1080.",
            "PASS" if scale_check_desktop.get("passed") else "FAIL",
            scale_check_desktop.get("reason", "Unknown extraction error"),
            full_ss,
        ))

        # 5. Color Contrast Validation (WCAG 2.1 / GIGW 3.0)
        contrast_check = check_color_contrast(driver)
        contrast_ss = take_full_screenshot(
            driver, os.path.join(ss_dir, "contrast_desktop_view.png")
        )
        results.append(AuditResult(
            "CONTRAST-01",
            "Accessible Color Pairing & Contrast",
            "Verify text/background combinations meet WCAG 2.1 minimum (4.5:1) standards.",
            "PASS" if contrast_check.get("passed") else "FAIL",
            contrast_check.get("reason", "Unknown extraction error"),
            contrast_ss,
        ))

        # 6. Type Scale Check (Mobile Viewport)
        driver.set_window_size(375, 812)
        time.sleep(1)
        mobile_ss = take_full_screenshot(driver, os.path.join(ss_dir, "mobile_view.png"))
        scale_check_mobile = check_type_scale(driver)
        results.append(AuditResult(
            "SCALE-02",
            "Typographic Hierarchy Scale (Mobile)",
            "Verify font sizes maintain a logical hierarchy (H1>H2>H3>P) at 375x812.",
            "PASS" if scale_check_mobile.get("passed") else "FAIL",
            scale_check_mobile.get("reason", "Unknown extraction error"),
            mobile_ss,
        ))

        # 7. Button Padding Consistency
        driver.set_window_size(1920, 1080)
        time.sleep(1)
        btn_padding = check_button_padding_consistency(driver)
        results.append(AuditResult(
            "BTN-01",
            "Button Padding Consistency",
            "Verify button sizes are consistent with uniform padding (DBIM 4.5).",
            "PASS" if btn_padding.get("passed") else "FAIL",
            btn_padding.get("reason", "Unknown extraction error"),
            full_ss,
        ))

        # 8. Button State Differentiation - hover / focus / disabled
        btn_states = check_button_states(driver)
        results.append(AuditResult(
            "BTN-02",
            "Button State Style Differentiation",
            "Verify distinct styles for enabled, hover, focus, disabled states (DBIM 4.5).",
            "PASS" if btn_states.get("passed") else "FAIL",
            btn_states.get("reason", "Unknown extraction error"),
            full_ss,
        ))

        # 9. Mouse Hover Feedback on Clickable Items
        hover_check = check_hover_feedback(driver)
        results.append(AuditResult(
            "BTN-03",
            "Mouse Hover Feedback on Clickable Items",
            "Verify mouse hover prompts a noticeable visual change on clickable items (DBIM 4.5).",
            "PASS" if hover_check.get("passed") else "FAIL",
            hover_check.get("reason", "Unknown extraction error"),
            full_ss,
        ))

        # -- Final Global Aggregation ------------------------------------------
        fail_count = sum(1 for r in results if r.result == "FAIL")
        overall    = "PASS" if fail_count == 0 else "FAIL"
        reason     = (f"Audit complete. Checked {len(results) - 1} rule definitions. "
                      f"Total failures: {fail_count}.")

        results.append(AuditResult(
            "SUMMARY",
            "Overall Compliance Verdict",
            "Aggregate summary across all typographic, structural, and content rules.",
            overall, reason, full_ss,
        ))

    finally:
        driver.quit()

    return results


# ---------------------------------------------------------------------------
# DOCX report generator
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn('w:shd'),
        {qn('w:fill'): color_hex, qn('w:val'): 'clear'}
    )
    shading.append(shd)


def apply_table_formatting(tbl, is_header_row_present=True):
    for row_idx, row in enumerate(tbl.rows):
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if is_header_row_present and row_idx == 0:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                scrubbed = re.sub(r'[\u20B9\$\%,\s]', '', text)
                try:
                    float(scrubbed)
                    is_num = len(scrubbed) > 0
                except ValueError:
                    is_num = False

                if is_num:
                    align = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    align = WD_ALIGN_PARAGRAPH.LEFT

            for p in cell.paragraphs:
                p.alignment = align


def generate_report(url: str, results: list, output_dir: str):
    doc = Document()

    style = doc.styles['Normal']
    font  = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # -- Title page -----------------------------------------------------------
    doc.add_paragraph()
    title = doc.add_heading("Typeface and Alignment Compliance Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Government of India - Digital Presence Guidelines Audit"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"URL: {url}\n").bold = True
    meta.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    meta.add_run("Tool: GIGW Auditor v2.6\n")
    doc.add_page_break()

    # -- Executive summary ----------------------------------------------------
    doc.add_heading("Executive Summary", level=1)
    overall = results[-1]
    color = (RGBColor(0x00, 0x80, 0x00)
             if overall.result == "PASS"
             else RGBColor(0xCC, 0x00, 0x00))
    p = doc.add_paragraph()
    p.add_run("Overall Result: ").bold = True
    r = p.add_run(overall.result)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(16)
    doc.add_paragraph(overall.reason)

    # -- Summary table --------------------------------------------------------
    doc.add_heading("Summary", level=2)
    pass_c = sum(1 for r in results if r.result == "PASS")
    fail_c = sum(1 for r in results if r.result == "FAIL")
    skip_c = sum(1 for r in results if r.result == "SKIP")
    info_c = sum(1 for r in results if r.result == "INFO")

    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Light List Accent 1'
    data = [
        ("Metric", "Count"),
        ("PASS",   str(pass_c)),
        ("FAIL",   str(fail_c)),
        ("SKIP",   str(skip_c)),
        ("INFO",   str(info_c)),
    ]
    for i, (k, v) in enumerate(data):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v

    apply_table_formatting(tbl, is_header_row_present=True)
    doc.add_page_break()

    # -- Detailed results -----------------------------------------------------
    doc.add_heading("Detailed Test Results", level=1)

    for res in results:
        doc.add_heading(f"{res.tc_id}: {res.name}", level=2)

        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        labels = ["Test Case", "Description", "Result", "Reason"]
        values = [res.tc_id, res.description, res.result, res.reason]
        colors = {
            "PASS":    "C6EFCE",
            "FAIL":    "FFC7CE",
            "SKIP":    "FFF2CC",
            "INFO":    "D9E2F3",
            "SUMMARY": "D9E2F3",
        }

        for i, (lbl, val) in enumerate(zip(labels, values)):
            tbl.rows[i].cells[0].text = lbl
            tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            tbl.rows[i].cells[1].text = val
            if lbl == "Result":
                set_cell_shading(tbl.rows[i].cells[1], colors.get(val, "FFFFFF"))

        apply_table_formatting(tbl, is_header_row_present=False)

        # ---------------- Screenshot embedding ----------------
        if res.screenshot and os.path.exists(res.screenshot):
            doc.add_paragraph()
            label_p = doc.add_paragraph()
            label_p.add_run("Screenshot Evidence:").bold = True

            embedded = False

            # First attempt: embed at 5.5 inches
            try:
                doc.add_picture(res.screenshot, width=Inches(5.5))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                embedded = True
            except Exception as e1:
                # Second attempt: downscale on-the-fly with Pillow and retry
                try:
                    from PIL import Image
                    tmp_path = res.screenshot.replace(".png", "_small.png")
                    img = Image.open(res.screenshot)
                    img.thumbnail((1400, 6000), Image.LANCZOS)
                    img.save(tmp_path, optimize=True)
                    img.close()
                    doc.add_picture(tmp_path, width=Inches(5.5))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    embedded = True
                except Exception as e2:
                    err_p = doc.add_paragraph(
                        f"[Could not embed image '{os.path.basename(res.screenshot)}': "
                        f"{type(e1).__name__}: {e1} | retry: {type(e2).__name__}: {e2}]"
                    )
                    err_p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

            # Always log the file path so user can inspect manually
            if embedded:
                path_p = doc.add_paragraph()
                run = path_p.add_run(f"(Source file: {res.screenshot})")
                run.italic = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        doc.add_paragraph()

    # -- Methodology ----------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "This audit uses Selenium WebDriver to load the target URL in a "
        "headless Chromium browser. For each HTML element type, the tool "
        "reads the computed font-family via window.getComputedStyle() and "
        "checks whether any Noto Sans variant is present. It also inspects "
        "<link> tags, @font-face rules, @import rules, and inline style "
        "attributes for references to Noto Sans. Screenshots are captured "
        "as evidence for each test case. Furthermore, it verifies text "
        "and table alignment using JavaScript to inspect computed styles. "
        "Content parsing evaluates paragraphs to ensure standard casing and "
        "prevents phonetic Hinglish terms. Finally, layout scripts confirm "
        "proportional descending header scales (Type Scale validation) and "
        "mathematically calculate precise RGB relative luminance algorithms to "
        "enforce GIGW 3.0 / WCAG 2.1 color contrast minimums."
    )

    doc.add_heading("Compliance Criteria", level=2)
    doc.add_paragraph(
        "As per the Government of India digital presence guidelines, "
        "Noto Sans must be selected as the main typeface. Furthermore, "
        "body text must be left-aligned, and tables must have left-aligned "
        "text, right-aligned numbers, and center-aligned column names. "
        "Long text blocks and paragraphs must not be fully capitalized, "
        "and content should avoid casual Hinglish transliteration. Type "
        "sizes must also remain visibly proportional descending from H1, "
        "and element background-to-text color contrast ratio must score "
        "at least 4.5:1 across the user interface."
    )

    # -- Save -----------------------------------------------------------------
    parsed    = urlparse(url)
    safe_name = re.sub(r'[^a-zA-Z0-9]', '_', parsed.netloc)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"Compliance_Audit_{safe_name}_{timestamp}.docx"
    filepath  = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Audit a website for Typeface and Alignment compliance."
    )
    parser.add_argument(
        "-o", "--output",
        default="audit_output",
        help="Output directory (default: audit_output)",
    )
    args = parser.parse_args()

    url = input("Enter Website URL: ").strip()
    if not url.startswith("http"):
        url = "https://" + url

    output_dir = args.output
    print("=" * 60)
    print("  Typeface & Alignment Compliance Auditor")
    print("  Gov. of India Digital Presence Guidelines")
    print("=" * 60)
    print(f"  Target : {url}")
    print(f"  Output : {os.path.abspath(output_dir)}")
    print("=" * 60)

    results = run_audit(url, output_dir)

    print("\n[*] Generating DOCX report ...")
    report_path = generate_report(url, results, output_dir)
    print(f"[+] Report saved: {report_path}")

    print("\n" + "-" * 60)
    icons = {"PASS": "OK", "FAIL": "FAIL", "SKIP": "SKIP", "INFO": "INFO"}
    for r in results:
        print(f"  [{icons.get(r.result, '?')}] {r.tc_id}: {r.name} -> {r.result}")
    print("-" * 60)

    fail_c = sum(1 for r in results if r.result == "FAIL")
    if fail_c:
        print(f"\n[!] {fail_c} functional criterion verification(s) FAILED.")
    else:
        print("\n[OK] All checks passed successfully - site structural layouts compliant!")

    return 0 if fail_c == 0 else 1


if __name__ == "__main__":
    main()