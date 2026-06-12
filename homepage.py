import os
import re
import sys
import time
import io                          # FIX 1: was "import ioimport requests" (missing newline)
import requests
from datetime import datetime
from urllib.parse import urljoin, urlparse

from PIL import Image
from bs4 import BeautifulSoup      # FIX 2: was "from bs4 import BeautifulSoupfrom docx..." (missing newline)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By

from webdriver_manager.chrome import ChromeDriverManager

try:
    import pytesseract              # FIX 3: was "import pytesseract    HAS_OCR" (missing newline)
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

OUTPUT_DIR = os.path.abspath("dbim_report")
os.makedirs(OUTPUT_DIR, exist_ok=True)

AUTHORIZED_PM_SOURCES = [
    "pmindia.gov.in", "pib.gov.in", "mygov.in", "india.gov.in"
]

SOCIAL_PLATFORMS = {
    "Twitter/X": ["twitter.com", "x.com"],
    "Instagram": ["instagram.com"],
    "Facebook": ["facebook.com"],
    "LinkedIn": ["linkedin.com"],
    "YouTube": ["youtube.com", "youtu.be"]
}

FOOTER_REQUIRED = [
    "archives", "website policies", "sitemap", "help", "feedback", "contact"
]


# =========================================================
# SETUP
# =========================================================

def setup_driver():
    options = Options()
    options.add_argument("--headless=new")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--window-size=1920,1080")
    options.add_argument(
        "user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )

    driver = webdriver.Chrome(
        service=Service(ChromeDriverManager().install()),
        options=options
    )

    driver.execute_cdp_cmd('Page.addScriptToEvaluateOnNewDocument', {
        'source': '''
            Object.defineProperty(navigator, "webdriver", {get: () => undefined});
            window.chrome = { runtime: {} };
        '''
    })

    return driver


def wait_for_full_load(driver, timeout=30, settle_time=8):
    try:
        WebDriverWait(driver, timeout).until(
            lambda d: d.execute_script("return document.readyState") == "complete"
        )
        time.sleep(settle_time)
    except:
        time.sleep(settle_time)


# =========================================================
# DOCUMENT SETUP
# =========================================================

doc = Document()

# Title
title = doc.add_heading("DBIM Compliance Report", level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run("Evidence Display Version v2.2")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

# Timestamp
p = doc.add_paragraph()
p.add_run("Generated: ").bold = True
p.add_run(datetime.now().strftime("%d %B %Y at %H:%M:%S"))

results = []


# =========================================================
# HELPER: ELEMENT TEXT & LINKS
# =========================================================

def get_element_text(driver, element):
    """Get rendered text from WebElement"""
    if not element:
        return ""
    try:
        return driver.execute_script("return arguments[0].innerText || '';", element)
    except:
        return ""


def get_element_links(driver, element):
    """Get all links in WebElement"""
    if not element:
        return []
    try:
        return driver.execute_script("""
            return [...arguments[0].querySelectorAll('a[href]')]
                .map(a => ({
                    text: a.innerText.trim(),
                    href: a.href
                }))
                .filter(item => item.href.startsWith('http'));
        """, element)
    except:
        return []


def get_all_rendered_links(driver):
    """
    Return every absolute URL that appears on the current page.
    """
    try:
        return driver.execute_script(r"""
            const anchors = [...document.querySelectorAll('a[href]')];
            const hrefs = anchors.map(a => a.href);
            return hrefs.filter(h => h && h.startsWith('http'));
        """)
    except Exception as e:
        print(f"⚠️  get_all_rendered_links failed: {e}")
        return []


# =========================================================
# PM CONTAINER DETECTION
# =========================================================

def find_pm_section_rendered(driver):
    """
    Locate the actual PM quote section on the homepage.
    """
    print("  🔍 Searching for PM section...")

    try:
        result = driver.execute_script(r"""
        function isVisible(el) {
            if (!el) return false;
            const style = window.getComputedStyle(el);
            return style.display !== 'none'
                && style.visibility !== 'hidden'
                && style.opacity !== '0'
                && el.offsetWidth > 0
                && el.offsetHeight > 0;
        }

        function scoreContainer(el) {
            const txt  = (el.innerText || '').toLowerCase();
            const html = el.innerHTML.toLowerCase();
            let score = 0;

            const pmKW = [
                'prime minister', 'narendra modi', "pm's quote",
                'hon\'ble prime', 'pradhan mantri', 'pm shri'
            ];
            for (const kw of pmKW) {
                if (txt.includes(kw)) { score += 15; break; }
            }
            if (score === 0) {
                const lightPM = ['modi ji', 'pm modi'];
                for (const kw of lightPM) {
                    if (txt.includes(kw)) { score += 5; break; }
                }
            }

            const eventKW = [
                'semicon', 'g20', 'summit', 'vibrant gujarat',
                'inaugurat', 'conference', 'ceremony', 'launch',
                'global', 'india 2025', 'international'
            ];
            let eventHit = false;
            for (const kw of eventKW) {
                if (txt.includes(kw)) { score += 10; eventHit = true; break; }
            }

            const datePat = /\b\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4}\b/;
            const datePat2 = /\b(?:\d{1,2}\s+)?(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}\b/i;
            if (datePat.test(txt) || datePat2.test(txt)) score += 8;

            const imgs = el.querySelectorAll('img');
            let hasPMImage = false;
            for (const img of imgs) {
                const alt = (img.alt || '').toLowerCase();
                const src = (img.src || '').toLowerCase();
                const ttl = (img.title || '').toLowerCase();
                if (alt.includes('modi') || alt.includes('prime minister')
                    || src.includes('modi') || src.includes('pm-')
                    || ttl.includes('modi')) {
                    hasPMImage = true;
                    break;
                }
            }
            if (hasPMImage) score += 12;

            const ctaKW = ['view event', 'read more', 'know more', 'learn more', 'view details'];
            const anchors = el.querySelectorAll('a[href]');
            let hasCTA = false;
            for (const a of anchors) {
                const aTxt = (a.innerText || '').toLowerCase();
                if (ctaKW.some(k => aTxt.includes(k))) { hasCTA = true; break; }
            }
            if (hasCTA) score += 8;

            const len = txt.length;
            if (len < 30) score -= 20;
            else if (len < 600) score += 5;
            else if (len < 2000) score += 2;
            else score -= 3;

            if (txt.includes('घोषणाएं') || txt.includes('announcement')) score -= 10;
            if (txt.includes('chevron_left') || txt.includes('chevron_right')) score -= 10;
            if (txt.includes('pause') && txt.includes('play')) score -= 10;

            return { el: el, score: score, eventHit: eventHit, hasPMImage: hasPMImage };
        }

        let bestResult = null;
        let bestScore = -999;

        const pmImgXPaths = [
            "//img[contains(translate(@alt,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'modi')]",
            "//img[contains(translate(@alt,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'prime minister')]",
            "//img[contains(translate(@src,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'modi')]",
            "//img[contains(translate(@title,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'modi')]"
        ];
        for (const xp of pmImgXPaths) {
            const result = document.evaluate(xp, document, null,
                XPathResult.ORDERED_NODE_SNAPSHOT_TYPE, null);
            for (let i = 0; i < result.snapshotLength; i++) {
                const img = result.snapshotItem(i);
                if (!isVisible(img)) continue;
                let container = img.parentElement;
                let steps = 0;
                while (container && steps < 15) {
                    steps++;
                    const s = scoreContainer(container);
                    if (s.score > bestScore) {
                        bestScore = s.score;
                        bestResult = container;
                    }
                    container = container.parentElement;
                }
            }
        }

        const ctaPhrases = ['view event', 'read more', 'know more', 'learn more', 'view details'];
        for (const phrase of ctaPhrases) {
            const links = document.querySelectorAll('a[href]');
            for (const link of links) {
                const lt = (link.innerText || '').toLowerCase();
                if (!lt.includes(phrase)) continue;
                if (!isVisible(link)) continue;
                let container = link.parentElement;
                let steps = 0;
                while (container && steps < 15) {
                    steps++;
                    const s = scoreContainer(container);
                    if (s.score > bestScore) {
                        bestScore = s.score;
                        bestResult = container;
                    }
                    container = container.parentElement;
                }
            }
        }

        const allBlocks = document.querySelectorAll('section, article, div[class], div[id]');
        for (const block of allBlocks) {
            if (!isVisible(block)) continue;
            const txt = (block.innerText || '');
            if (txt.length < 40 || txt.length > 5000) continue;
            const s = scoreContainer(block);
            if (s.score > bestScore) {
                bestScore = s.score;
                bestResult = block;
            }
        }

        if (bestResult) {
            const finalScore = scoreContainer(bestResult);
            if (finalScore.score >= 5) {
                return {
                    found: true,
                    score: finalScore.score,
                    hasPMImage: finalScore.hasPMImage,
                    eventHit: finalScore.eventHit,
                    textPreview: (bestResult.innerText || '').substring(0, 300),
                    tag: bestResult.tagName,
                    className: (bestResult.className || '').substring(0, 100)
                };
            }
        }

        return { found: false };
        """)

        if result and result.get('found'):
            container = _relocate_pm_container(driver, result)
            if container:
                print(f"  ✅ PM container found:")
                print(f"      Score   : {result.get('score', '?')}/35")
                print(f"      PM image: {result.get('hasPMImage', False)}")
                print(f"      Event   : {result.get('eventHit', False)}")
                print(f"      Tag     : {result.get('tag', '?')}")
                print(f"      Class   : {result.get('className', '?')[:60]}")
                print(f"      Preview : {result.get('textPreview', '')[:120]}...")
                return container

        print("  ❌ PM container not found (all strategies exhausted)")
        return None

    except Exception as e:
        print(f"  ❌ PM search error: {e}")
        import traceback
        traceback.print_exc()
        return None


def _relocate_pm_container(driver, js_result):
    """
    Given diagnostic info from the JS scoring pass, locate the actual
    Selenium WebElement for the winning container.
    """
    try:
        container = driver.execute_script(r"""
            function isVisible(el) {
                if (!el) return false;
                const s = window.getComputedStyle(el);
                return s.display !== 'none' && s.visibility !== 'hidden'
                    && el.offsetWidth > 0 && el.offsetHeight > 0;
            }

            function scoreEl(el) {
                const txt = (el.innerText || '').toLowerCase();
                let score = 0;
                const pmKW = ['prime minister','narendra modi',"pm's quote",
                              'hon\'ble prime','pradhan mantri','pm shri'];
                for (const kw of pmKW) { if (txt.includes(kw)) { score += 15; break; } }
                if (score === 0) {
                    for (const kw of ['modi ji','pm modi']) { if (txt.includes(kw)) { score += 5; break; } }
                }
                const evKW = ['semicon','g20','summit','vibrant gujarat',
                              'inaugurat','conference','ceremony','launch',
                              'global','india 2025','international'];
                let evHit = false;
                for (const kw of evKW) { if (txt.includes(kw)) { score += 10; evHit = true; break; } }
                if (/\b\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4}\b/.test(txt)) score += 8;
                if (/\b(?:\d{1,2}\s+)?(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}\b/i.test(txt)) score += 8;
                const imgs = el.querySelectorAll('img');
                let pmImg = false;
                for (const img of imgs) {
                    const a = (img.alt||'').toLowerCase(), s2 = (img.src||'').toLowerCase();
                    if (a.includes('modi')||a.includes('prime minister')||s2.includes('modi')) { pmImg=true; break; }
                }
                if (pmImg) score += 12;
                const ctaKW = ['view event','read more','know more','learn more','view details'];
                for (const a of el.querySelectorAll('a[href]')) {
                    if (ctaKW.some(k => (a.innerText||'').toLowerCase().includes(k))) { score += 8; break; }
                }
                const len = txt.length;
                if (len < 30) score -= 20;
                else if (len < 600) score += 5;
                else if (len < 2000) score += 2;
                else score -= 3;
                if (txt.includes('घोषणा') || txt.includes('chevron_left')) score -= 10;
                return score;
            }

            let bestEl = null, bestScore = -999;

            const xps = [
                "//img[contains(translate(@alt,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'modi')]",
                "//img[contains(translate(@alt,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'prime minister')]",
                "//img[contains(translate(@src,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'modi')]"
            ];
            for (const xp of xps) {
                const r = document.evaluate(xp, document, null, XPathResult.ORDERED_NODE_SNAPSHOT_TYPE, null);
                for (let i = 0; i < r.snapshotLength; i++) {
                    const img = r.snapshotItem(i);
                    if (!isVisible(img)) continue;
                    let c = img.parentElement, steps = 0;
                    while (c && steps < 15) {
                        steps++;
                        const sc = scoreEl(c);
                        if (sc > bestScore) { bestScore = sc; bestEl = c; }
                        c = c.parentElement;
                    }
                }
            }

            const cta = ['view event','read more','know more','learn more','view details'];
            for (const phrase of cta) {
                for (const a of document.querySelectorAll('a[href]')) {
                    if (!(a.innerText||'').toLowerCase().includes(phrase)) continue;
                    if (!isVisible(a)) continue;
                    let c = a.parentElement, steps = 0;
                    while (c && steps < 15) {
                        steps++;
                        const sc = scoreEl(c);
                        if (sc > bestScore) { bestScore = sc; bestEl = c; }
                        c = c.parentElement;
                    }
                }
            }

            for (const block of document.querySelectorAll('section, article, div[class], div[id]')) {
                if (!isVisible(block)) continue;
                const t = (block.innerText || '');
                if (t.length < 40 || t.length > 5000) continue;
                const sc = scoreEl(block);
                if (sc > bestScore) { bestScore = sc; bestEl = block; }
            }

            return (bestEl && bestScore >= 5) ? bestEl : null;
        """)

        return container
    except Exception as e:
        print(f"  ⚠️  Re-location failed: {e}")
        return None


# =========================================================
# SCREENSHOT UTILITIES
# =========================================================

def save_screenshot(driver, name):
    path = os.path.join(OUTPUT_DIR, f"{name}.png")
    try:
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)
        driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(1)

        total_width = driver.execute_script(
            "return Math.max(document.documentElement.scrollWidth, document.body.scrollWidth);"
        )
        total_height = driver.execute_script(
            "return Math.max(document.documentElement.scrollHeight, document.body.scrollHeight);"
        )

        driver.set_window_size(max(total_width, 1920), min(total_height + 200, 25000))
        time.sleep(1)
        driver.save_screenshot(path)

        if os.path.exists(path) and os.path.getsize(path) > 15000:
            print(f"  📸 {name}.png ({os.path.getsize(path)//1024} KB)")
            return path

    except Exception as e:
        print(f"  ⚠️  Screenshot failed: {e}")

    return None


# =========================================================
# ENHANCED REPORT UTILITIES WITH EVIDENCE DISPLAY
# =========================================================

def add_result(test_id, title, status, reason, evidence="", screenshot=None):
    """Add test result with FULL evidence display"""
    results.append({
        "id": test_id, "title": title, "status": status,
        "reason": reason, "evidence": evidence, "screenshot": screenshot
    })

    icons = {"PASS": "✅", "FAIL": "❌", "PARTIAL": "⚠️", "MANUAL_CHECK": "🔍"}
    icon = icons.get(status, "•")

    print(f"\n  {icon} [{test_id}] {status}")
    print(f"      Finding: {reason}")
    if evidence:
        print(f"      Evidence: {evidence}")

    if len(results) > 1:
        doc.add_page_break()

    heading = doc.add_heading(f"{test_id}: {title}", level=1)

    p = doc.add_paragraph()
    p.add_run("Status: ").bold = True
    sr = p.add_run(status)
    sr.bold = True
    sr.font.size = Pt(12)

    colors = {
        "PASS": RGBColor(0, 128, 0),
        "FAIL": RGBColor(200, 0, 0),
        "PARTIAL": RGBColor(255, 140, 0),
        "MANUAL_CHECK": RGBColor(100, 100, 200)
    }
    sr.font.color.rgb = colors.get(status, RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.add_run("Finding: ").bold = True
    p.add_run(reason)

    if evidence:
        p = doc.add_paragraph()
        p.add_run("Evidence: ").bold = True
        if " | " in evidence:
            doc.add_paragraph()
            parts = evidence.split(" | ")
            for part in parts:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.add_run(part)
        else:
            p.add_run(evidence)

    if screenshot and os.path.exists(screenshot):
        doc.add_paragraph()
        doc.add_paragraph().add_run("Screenshot:").bold = True
        try:
            doc.add_picture(screenshot, width=Inches(6.5))
        except Exception as e:
            doc.add_paragraph(f"[Screenshot error: {e}]")


def get_soup(driver):
    return BeautifulSoup(driver.page_source, "html.parser")


# =========================================================
# TEST FUNCTIONS
# =========================================================

def test_01_header(driver, soup):
    """DBIM-01: Header with Logo"""
    screenshot = save_screenshot(driver, "01_header")

    header = soup.find("header") or soup.find("nav") or soup.find(class_=re.compile(r"header|navbar", re.I))
    logo = soup.find("img", alt=re.compile(r"logo|emblem|ashoka", re.I))
    search = soup.find("input", {"type": "search"}) or soup.find(class_=re.compile(r"search", re.I))

    evidence_parts = []
    evidence_parts.append(f"Header element: {'Found' if header else 'Not found'}")
    evidence_parts.append(f"Logo/Emblem: {'Found' if logo else 'Not found'}")
    evidence_parts.append(f"Search box: {'Found' if search else 'Not found'}")

    if logo:
        evidence_parts.append(f"Logo alt text: {logo.get('alt', 'N/A')[:50]}")

    evidence = " | ".join(evidence_parts)

    if header and logo:
        add_result("DBIM-01", "Header with Logo (Section i)", "PASS",
                   "Header contains logo/emblem.", evidence, screenshot)
    else:
        add_result("DBIM-01", "Header with Logo (Section i)", "FAIL",
                   "Header or logo missing.", evidence, screenshot)


def test_02_breadcrumbs(driver, soup):
    """DBIM-02: Breadcrumb Navigation"""
    screenshot = save_screenshot(driver, "02_breadcrumbs")

    breadcrumb = (
        soup.find(attrs={"aria-label": "breadcrumb"}) or
        soup.find(class_=re.compile(r"breadcrumb", re.I)) or
        soup.find("nav", class_=re.compile(r"breadcrumb", re.I))
    )

    evidence = f"Breadcrumb element: {'Found' if breadcrumb else 'Not found'}"
    if breadcrumb:
        evidence += f" | Tag: {breadcrumb.name} | Class: {breadcrumb.get('class', 'N/A')}"

    if breadcrumb:
        add_result("DBIM-02", "Breadcrumb Navigation (Section ii)", "PASS",
                   "Breadcrumb navigation found.", evidence, screenshot)
    else:
        add_result("DBIM-02", "Breadcrumb Navigation (Section ii)", "FAIL",
                   "Breadcrumb navigation missing.", evidence, screenshot)


def test_03_carousel(driver, soup):
    """DBIM-03: Auto-Rotating Carousel"""
    screenshot = save_screenshot(driver, "03_carousel")

    carousel = (
        soup.find(class_=re.compile(r"carousel|slider|slideshow|swiper", re.I)) or
        soup.find(id=re.compile(r"carousel|slider|slideshow", re.I))
    )

    evidence = f"Carousel element: {'Found' if carousel else 'Not found'}"
    if carousel:
        evidence += f" | Tag: {carousel.name} | Class: {carousel.get('class', 'N/A')[:50]}"

    if carousel:
        add_result("DBIM-03", "Auto-Rotating Carousel (Section iii)", "PASS",
                   "Carousel/slider found on homepage.", evidence, screenshot)
    else:
        add_result("DBIM-03", "Auto-Rotating Carousel (Section iii)", "FAIL",
                   "No carousel/slider detected.", evidence, screenshot)


def test_05_pm_sources(driver):
    """DBIM-05: Authorized PM Sources"""
    screenshot = save_screenshot(driver, "05_pm_sources")

    all_links = get_all_rendered_links(driver)

    authorized = []
    for link in all_links:
        for src in AUTHORIZED_PM_SOURCES:
            if src in link.lower():
                authorized.append(link)
                break
    authorized = list(set(authorized))[:5]

    evidence_parts = []
    evidence_parts.append(f"Total links checked: {len(all_links)}")
    evidence_parts.append(f"Authorized sources found: {len(authorized)}")

    if authorized:
        evidence_parts.append(f"Required sources: {', '.join(AUTHORIZED_PM_SOURCES)}")
        evidence_parts.append("Links found:")
        for i, link in enumerate(authorized[:3], 1):
            evidence_parts.append(f"  {i}. {link[:80]}")

    evidence = " | ".join(evidence_parts)

    if authorized:
        add_result("DBIM-05", "PM Content from Authorized Sources (Section iv)", "PASS",
                   f"Found {len(authorized)} authorized PM source link(s).",
                   evidence, screenshot)
    else:
        add_result("DBIM-05", "PM Content from Authorized Sources (Section iv)", "FAIL",
                   "No authorized PM sources found. Required: pmindia.gov.in, pib.gov.in, mygov.in",
                   evidence, screenshot)


def test_06_pm_quote_date(driver, pm_container):
    """DBIM-06: PM Quote Date & Relevance"""
    screenshot = save_screenshot(driver, "06_pm_quote_date")

    if not pm_container:
        add_result("DBIM-06", "PM Quote Relevance & Date (Section iv)", "FAIL",
                   "PM section not found on page.",
                   "Container detection: Failed | Text extracted: 0 chars",
                   screenshot)
        return

    text = get_element_text(driver, pm_container)
    text_length = len(text)

    print(f"  📝 Container text: {text_length} chars")
    print(f"  📝 Preview: {text[:250]}")

    if text_length < 20:
        evidence = f"Container found: Yes | Text length: {text_length} chars | Status: Insufficient text"
        add_result("DBIM-06", "PM Quote Relevance & Date (Section iv)", "FAIL",
                   "PM container has insufficient extractable text.",
                   evidence, screenshot)
        return

    date_patterns = [
        (r"\b\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4}\b", "DD.MM.YYYY"),
        (r"\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}\b", "DD Mon YYYY"),
        (r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{1,2},?\s+\d{4}\b", "Mon DD YYYY"),
        (r"\b\d{4}[.\-/]\d{1,2}[.\-/]\d{1,2}\b", "YYYY-MM-DD"),
    ]

    date_found = None
    date_format = None

    for pattern, fmt in date_patterns:
        match = re.search(pattern, text, re.I)
        if match:
            date_found = match.group()
            date_format = fmt
            print(f"  ✅ Date: {date_found} ({fmt})")
            break

    event_keywords = {
        'high': ['semicon', 'g20', 'summit', 'vibrant gujarat'],
        'medium': ['conference', 'inauguration', 'launch', 'ceremony'],
        'low': ['event', 'address', 'speech']
    }

    event_score = 0
    event_name = None
    event_priority = None

    text_lower = text.lower()

    for priority, keywords in event_keywords.items():
        for kw in keywords:
            if kw in text_lower:
                if priority == 'high':
                    event_score = 5
                    event_priority = 'High'
                elif priority == 'medium' and event_score < 5:
                    event_score = 3
                    event_priority = 'Medium'
                elif priority == 'low' and event_score < 3:
                    event_score = 1
                    event_priority = 'Low'

                lines = text.split('\n')
                for line in lines:
                    if kw in line.lower() and 5 < len(line.strip()) < 100:
                        event_name = line.strip()
                        break

                if event_score >= 3:
                    break

        if event_score >= 3:
            break

    print(f"  🎯 Event: {event_name or 'Not found'} (score: {event_score}/5, priority: {event_priority or 'None'})")

    evidence_parts = []
    evidence_parts.append(f"Text extracted: {text_length} characters")
    evidence_parts.append(f"Text preview: {text[:150].strip()}")
    evidence_parts.append(f"Date found: {date_found or 'Not detected'}")
    if date_format:
        evidence_parts.append(f"Date format: {date_format}")
    evidence_parts.append(f"Event detected: {event_name or 'Not detected'}")
    evidence_parts.append(f"Event confidence: {event_score}/5 ({event_priority or 'None'})")
    evidence_parts.append(f"Keywords searched: PM, Event, Date patterns")

    evidence = " | ".join(evidence_parts)

    if date_found and event_score >= 3:
        add_result("DBIM-06", "PM Quote Relevance & Date (Section iv)", "PASS",
                   f"PM quote contains date ({date_found}) and event reference ({event_name}).",
                   evidence, screenshot)
    elif date_found and event_score >= 1:
        add_result("DBIM-06", "PM Quote Relevance & Date (Section iv)", "PASS",
                   f"PM quote contains date ({date_found}) with event context.",
                   evidence, screenshot)
    elif date_found:
        add_result("DBIM-06", "PM Quote Relevance & Date (Section iv)", "PARTIAL",
                   f"Date found ({date_found}) but event context weak.",
                   evidence, screenshot)
    elif event_score >= 3:
        add_result("DBIM-06", "PM Quote Relevance & Date (Section iv)", "PARTIAL",
                   f"Strong event reference ({event_name}) but no explicit date.",
                   evidence, screenshot)
    else:
        add_result("DBIM-06", "PM Quote Relevance & Date (Section iv)", "FAIL",
                   "Neither date nor strong event reference found in PM section.",
                   evidence, screenshot)


def test_07_pm_quote_hyperlink(driver, pm_container):
    """DBIM-07: PM Quote Hyperlink"""
    screenshot = save_screenshot(driver, "07_pm_quote_hyperlink")

    if not pm_container:
        add_result("DBIM-07", "PM Quote Format & Hyperlink (Section iv)", "FAIL",
                   "PM section not found.",
                   "Container detection: Failed",
                   screenshot)
        return

    links = get_element_links(driver, pm_container)

    event_keywords = ['view event', 'read more', 'know more', 'learn more', 'details', 'view quote']
    event_ctas = [
        link for link in links
        if any(kw in link['text'].lower() for kw in event_keywords)
    ]

    pm_domain_links = [
        link for link in links
        if any(domain in link['href'].lower() for domain in AUTHORIZED_PM_SOURCES)
    ]

    print(f"  🔗 Links in container: {len(links)}")
    print(f"  📌 Event CTAs: {len(event_ctas)}")
    print(f"  🏛️  PM domain links: {len(pm_domain_links)}")

    if event_ctas:
        for cta in event_ctas[:2]:
            print(f"      • {cta['text']} → {cta['href'][:60]}")

    evidence_parts = []
    evidence_parts.append(f"Total links in PM section: {len(links)}")
    evidence_parts.append(f"Event CTA buttons: {len(event_ctas)}")
    evidence_parts.append(f"PM domain links: {len(pm_domain_links)}")
    evidence_parts.append(f"Keywords searched: {', '.join(event_keywords[:4])}")

    if event_ctas:
        evidence_parts.append("Event CTAs found:")
        for i, cta in enumerate(event_ctas[:3], 1):
            evidence_parts.append(f"  {i}. '{cta['text']}' → {cta['href'][:60]}")

    if pm_domain_links:
        evidence_parts.append("PM domain links:")
        for i, link in enumerate(pm_domain_links[:2], 1):
            evidence_parts.append(f"  {i}. {link['href'][:60]}")

    if not links:
        evidence_parts.append("No links found in PM section")

    evidence = " | ".join(evidence_parts)

    if event_ctas:
        add_result("DBIM-07", "PM Quote Format & Hyperlink (Section iv)", "PASS",
                   f"Found {len(event_ctas)} event CTA button(s) with hyperlink(s).",
                   evidence, screenshot)
    elif pm_domain_links:
        add_result("DBIM-07", "PM Quote Format & Hyperlink (Section iv)", "PASS",
                   f"Found {len(pm_domain_links)} PM domain link(s).",
                   evidence, screenshot)
    elif links:
        add_result("DBIM-07", "PM Quote Format & Hyperlink (Section iv)", "PARTIAL",
                   f"Found {len(links)} link(s) but no clear event CTA.",
                   evidence, screenshot)
    else:
        add_result("DBIM-07", "PM Quote Format & Hyperlink (Section iv)", "FAIL",
                   "No hyperlinks or buttons found in PM section.",
                   evidence, screenshot)


def test_14_social_media(driver):
    """DBIM-14: Social Media Integration"""
    screenshot = save_screenshot(driver, "14_social_media")

    all_links = get_all_rendered_links(driver)

    found_platforms = {}
    platform_links = {}

    for platform, domains in SOCIAL_PLATFORMS.items():
        found = False
        for link in all_links:
            if any(domain in link.lower() for domain in domains):
                found = True
                platform_links[platform] = link
                break
        found_platforms[platform] = found

    present = [k for k, v in found_platforms.items() if v]
    missing = [k for k, v in found_platforms.items() if not v]

    evidence_parts = []
    evidence_parts.append(f"Total links scanned: {len(all_links)}")
    evidence_parts.append(f"Platforms found: {len(present)}/5")

    if present:
        evidence_parts.append(f"Present: {', '.join(present)}")
        for platform in present:
            if platform in platform_links:
                evidence_parts.append(f"{platform}: {platform_links[platform][:60]}")

    if missing:
        evidence_parts.append(f"Missing: {', '.join(missing)}")

    evidence = " | ".join(evidence_parts)

    if len(missing) == 0:
        add_result("DBIM-14", "Social Media Integration (Section xi)", "PASS",
                   "All 5 required social media platforms found.", evidence, screenshot)
    elif len(present) >= 3:
        add_result("DBIM-14", "Social Media Integration (Section xi)", "PARTIAL",
                   f"{len(present)}/5 platforms found. Missing: {', '.join(missing)}",
                   evidence, screenshot)
    else:
        add_result("DBIM-14", "Social Media Integration (Section xi)", "FAIL",
                   f"Only {len(present)}/5 platforms found. Insufficient coverage.",
                   evidence, screenshot)


def test_15_footer(driver, soup):
    """
    DBIM-15: Footer with Links (Section xii)

    ROOT CAUSE FIX: The footer content on this site is injected by JavaScript
    *after* the initial HTML response.  BeautifulSoup only parses the raw HTML
    snapshot (driver.page_source), so it sees an empty footer skeleton.

    Fix: use Selenium's execute_script to read the *live rendered* DOM text,
    which reflects whatever JavaScript has injected.  BeautifulSoup is kept as
    a fallback detector only (to confirm a footer element exists).
    """
    screenshot = save_screenshot(driver, "15_footer")

    # ------------------------------------------------------------------
    # STEP 1 – Confirm a footer element is present in the live DOM
    # ------------------------------------------------------------------
    footer_info = driver.execute_script(r"""
        // Try <footer> first, then any element whose class contains "footer"
        let el = document.querySelector('footer');
        if (!el) {
            const all = document.querySelectorAll('[class*="footer"],[id*="footer"]');
            for (const candidate of all) {
                if (candidate.offsetWidth > 0 && candidate.offsetHeight > 0) {
                    el = candidate;
                    break;
                }
            }
        }
        if (!el) return { found: false };

        // Scroll it into view so it fully renders
        el.scrollIntoView({ behavior: 'instant', block: 'end' });

        return {
            found: true,
            tag: el.tagName,
            className: (el.className || '').substring(0, 80),
            // innerText gives the *rendered* text including JS-injected content
            text: (el.innerText || '').toLowerCase(),
            // Also grab every anchor href for bonus evidence
            links: [...el.querySelectorAll('a[href]')].map(a => ({
                text: (a.innerText || '').trim().toLowerCase(),
                href: a.href
            }))
        };
    """)

    if not footer_info or not footer_info.get('found'):
        # Last-resort: try BeautifulSoup (catches SSR-rendered footers)
        soup_footer = soup.find("footer") or soup.find(class_=re.compile(r"footer", re.I))
        if not soup_footer:
            add_result(
                "DBIM-15",
                "Footer with Links (Section xii)",
                "FAIL",
                "Footer element not found in live DOM or static HTML.",
                "Footer tag: Not found | Selenium: Not found | BeautifulSoup: Not found",
                screenshot,
            )
            return

        # BeautifulSoup fallback: give it one more second and re-read
        time.sleep(2)
        footer_info = driver.execute_script(r"""
            let el = document.querySelector('footer');
            if (!el) {
                const all = document.querySelectorAll('[class*="footer"],[id*="footer"]');
                for (const c of all) { if (c.offsetWidth > 0) { el = c; break; } }
            }
            return el ? { found: true, tag: el.tagName,
                          className: (el.className||'').substring(0,80),
                          text: (el.innerText||'').toLowerCase(),
                          links: [...el.querySelectorAll('a[href]')].map(a=>({
                              text:(a.innerText||'').trim().toLowerCase(), href:a.href}))
                        } : { found: false };
        """)
        if not footer_info or not footer_info.get('found'):
            add_result(
                "DBIM-15",
                "Footer with Links (Section xii)",
                "FAIL",
                "Footer element not found.",
                "Footer tag: Not found",
                screenshot,
            )
            return

    footer_text  = footer_info.get('text', '')
    footer_links = footer_info.get('links', [])

    print(f"  🦶 Footer tag   : {footer_info.get('tag')} .{footer_info.get('className')[:50]}")
    print(f"  🦶 Footer text  : {len(footer_text)} chars")
    print(f"  🦶 Footer links : {len(footer_links)}")
    print(f"  🦶 Preview      : {footer_text[:300]}")
    print(f"  🦶 Anchor texts : {[lnk['text'] for lnk in footer_links]}")

    # ------------------------------------------------------------------
    # STEP 2 – Normalise all text for robust matching
    #
    # Problems seen in the wild on this site:
    #   • Links are prefixed with ">" or "›" chevrons  (e.g. "> Archives")
    #   • "Contact" is labelled "Contact Us" — substring match handles this
    #   • Non-breaking spaces (\xa0) and Unicode whitespace fool \b boundaries
    #   • innerText may return mixed-case text
    # Fix: strip every non-alpha-numeric character and use plain substring
    # matching on the cleaned text instead of word-boundary regex.
    # ------------------------------------------------------------------

    def normalise(text):
        """Lowercase, collapse all whitespace/punctuation to single spaces."""
        text = text.replace('\xa0', ' ').replace('\u200b', '')
        text = re.sub(r'^[\s>›»\u2019•\-\u2013\u2014]+', '', text.strip())
        text = re.sub(r'\s+', ' ', text).strip().lower()
        return text

    # Build candidate list: full block + each anchor text, all normalised
    candidate_texts = [normalise(footer_text)] + [normalise(lnk['text']) for lnk in footer_links]

    print(f"  🦶 Normalised anchors: {[normalise(lnk['text']) for lnk in footer_links]}")

    # Each keyword maps to English + Hindi surface forms.
    # Hindi translations observed live on this site are listed as aliases.
    # A keyword PASSES if ANY alias appears as a substring in ANY candidate.
    required_keywords_map = {
        "archives":         [
            "archives", "archive",
            "अभिलेखागार", "अभिलेख",
        ],
        "website policies": [
            "website policies", "web policies", "site policies",
            "वेबसाइट नीतियाँ", "वेबसाइट नीति", "नीतियाँ",
        ],
        "sitemap":          [
            "sitemap", "site map",
            "साइटमैप",
        ],
        "help":             [
            "help",
            "मदद", "सहायता",
        ],
        "feedback":         [
            "feedback",
            "फ़ीडबैक", "फीडबैक", "प्रतिक्रिया",
        ],
        "contact":          [
            "contact",
            "संपर्क", "हमसे संपर्क",
        ],
    }

    found_items   = []
    missing_items = []

    for kw_label, aliases in required_keywords_map.items():
        matched = any(
            alias in candidate
            for candidate in candidate_texts
            for alias in aliases
        )
        if matched:
            found_items.append(kw_label)
        else:
            missing_items.append(kw_label)

    total_required = len(required_keywords_map)
    coverage = len(found_items) / total_required * 100

    # ------------------------------------------------------------------
    # STEP 3 – Build evidence string
    # ------------------------------------------------------------------
    evidence_parts = [
        f"Footer element: Found ({footer_info.get('tag')})",
        f"Detection method: Selenium live DOM (JS-rendered)",
        f"Footer text length: {len(footer_text)} chars",
        f"Footer anchor count: {len(footer_links)}",
        f"Required items: {total_required}",
        f"Items found: {len(found_items)} ({coverage:.0f}%)",
    ]
    if found_items:
        evidence_parts.append(f"Present: {', '.join(found_items)}")
    if missing_items:
        evidence_parts.append(f"Missing: {', '.join(missing_items)}")

    evidence = " | ".join(evidence_parts)

    # ------------------------------------------------------------------
    # STEP 4 – PASS / PARTIAL / FAIL
    # ------------------------------------------------------------------
    if coverage >= 80:
        result_status = "PASS"
        result_msg = f"Footer contains {len(found_items)}/{total_required} required items."
    elif coverage >= 50:
        result_status = "PARTIAL"
        result_msg = f"Footer incomplete. Missing: {', '.join(missing_items)}"
    else:
        result_status = "FAIL"
        result_msg = "Footer missing multiple required elements."


    add_result(
        "DBIM-15",
        "Footer with Links (Section xii)",
        result_status,
        result_msg,
        evidence,
        screenshot,
    )


# =========================================================
# SUMMARY
# =========================================================

def add_summary():
    """Add summary with full evidence table"""
    doc.add_page_break()

    heading = doc.add_heading("Compliance Summary", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    total = len(results)
    passed = sum(1 for r in results if r["status"] == "PASS")
    failed = sum(1 for r in results if r["status"] == "FAIL")
    partial = sum(1 for r in results if r["status"] in ["PARTIAL", "MANUAL_CHECK"])

    compliance_rate = (passed / total * 100) if total > 0 else 0

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run(f"Overall Compliance: {compliance_rate:.1f}%")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Total Tests: {total}\n").bold = True

    pr = p.add_run(f"✓ Passed: {passed}\n")
    pr.font.color.rgb = RGBColor(0, 128, 0)
    pr.font.size = Pt(12)

    fr = p.add_run(f"✗ Failed: {failed}\n")
    fr.font.color.rgb = RGBColor(200, 0, 0)
    fr.font.size = Pt(12)

    ppr = p.add_run(f"⚠ Partial: {partial}")
    ppr.font.color.rgb = RGBColor(255, 140, 0)
    ppr.font.size = Pt(12)

    doc.add_page_break()
    doc.add_heading("Detailed Test Results", level=2)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ["Test ID", "Section", "Status", "Finding", "Evidence Summary"]

    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for result in results:
        row_cells = table.add_row().cells
        row_cells[0].text = result["id"]
        row_cells[1].text = result["title"][:40]
        row_cells[2].text = result["status"]
        row_cells[3].text = result["reason"][:100]

        evidence = result.get("evidence", "")
        if evidence:
            evidence_summary = evidence.split("|")[0][:80]
            row_cells[4].text = evidence_summary
        else:
            row_cells[4].text = "N/A"


# =========================================================
# MAIN EXECUTION
# =========================================================

def main():
    global driver

    url = sys.argv[1] if len(sys.argv) > 1 else input("🌐 Enter website URL: ").strip()
    if not url.startswith("http"):
        url = "https://" + url

    print("\n" + "=" * 70)
    print("DBIM COMPLIANCE CHECKER v2.2 - EVIDENCE DISPLAY")
    print("=" * 70)
    print(f"Target: {url}")
    print(f"Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 70 + "\n")

    p = doc.add_paragraph()
    p.add_run("Target URL: ").bold = True
    p.add_run(url)

    driver = setup_driver()

    try:
        print("🌐 Loading page...")
        driver.get(url)
        wait_for_full_load(driver)

        print("\n📸 Capturing homepage...")
        homepage_ss = save_screenshot(driver, "00_homepage")

        if homepage_ss:
            doc.add_page_break()
            doc.add_heading("Homepage Screenshot", level=1)
            doc.add_picture(homepage_ss, width=Inches(6.5))

        soup = get_soup(driver)

        print("\n🔍 Locating PM section...")
        pm_container = find_pm_section_rendered(driver)

        print("\n" + "=" * 70)
        print("RUNNING COMPLIANCE TESTS")
        print("=" * 70)

        print("\n[01] Header with Logo...")
        test_01_header(driver, soup)

        print("\n[02] Breadcrumb Navigation...")
        test_02_breadcrumbs(driver, soup)

        print("\n[03] Auto-Rotating Carousel...")
        test_03_carousel(driver, soup)

        print("\n[05] PM Authorized Sources...")
        test_05_pm_sources(driver)

        print("\n[06] PM Quote Date & Relevance...")
        test_06_pm_quote_date(driver, pm_container)

        print("\n[07] PM Quote Hyperlink...")
        test_07_pm_quote_hyperlink(driver, pm_container)

        print("\n[14] Social Media Integration...")
        test_14_social_media(driver)

        print("\n[15] Footer Elements...")
        test_15_footer(driver, soup)

        add_summary()

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        report_filename = f"DBIM_Report_v2.2_{timestamp}.docx"
        report_path = os.path.join(OUTPUT_DIR, report_filename)
        doc.save(report_path)

        passed = sum(1 for r in results if r["status"] == "PASS")
        failed = sum(1 for r in results if r["status"] == "FAIL")
        partial = sum(1 for r in results if r["status"] in ["PARTIAL", "MANUAL_CHECK"])
        total = len(results)

        compliance = (passed / total * 100) if total > 0 else 0

        print("\n" + "=" * 70)
        print("✅ COMPLIANCE AUDIT COMPLETE")
        print("=" * 70)
        print(f"📄 Report: {report_path}")
        print(f"📊 Results: {passed} passed | {failed} failed | {partial} partial")
        print(f"📈 Compliance Rate: {compliance:.1f}%")
        print("=" * 70 + "\n")

    except Exception as e:
        print(f"\n❌ CRITICAL ERROR: {e}")
        import traceback
        traceback.print_exc()

    finally:
        driver.quit()
        print("🔒 Browser closed\n")


if __name__ == "__main__":
    main()